"""Template management module for weekly report agent.

Provides functions to list, get, save, delete, and extract templates.
Supports both builtin and user-defined custom templates.
"""

import json
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

_TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"
_BUILTIN_DIR = _TEMPLATES_DIR / "builtin"
_CUSTOM_DIR = _TEMPLATES_DIR / "custom"

# Template JSON schema keys
_REQUIRED_KEYS = {"name", "description", "template", "sections", "created_at"}


# ---------------------------------------------------------------------------
# File reading helpers
# ---------------------------------------------------------------------------

def _read_docx(file_path: Path) -> str:
    """Read a .docx file using python-docx. Returns markdown-like formatted content."""
    try:
        from docx import Document
        doc = Document(str(file_path))
        
        lines = []
        
        # Common heading keywords for detection
        heading_keywords = [
            '工作总结', '工作汇总', '本周工作', '近期工作',
            '重点工作', '项目进展', '工作内容',
            '遇到的问题', '问题与解决', '问题', '风险',
            '后续计划', '下周计划', '计划', '下一步',
            '完成事项', '已完成', '完成',
            '进行中', '进行中的工作',
            '需要的支持', '需要协助',
            '技术决策', '决策',
        ]
        
        # Read paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append('')
                continue
            
            # Detect heading styles - safely check style
            try:
                style_name = para.style.name if para.style else ''
            except Exception:
                style_name = ''
            
            # Check if it's a heading style
            if style_name and style_name.startswith('Heading'):
                try:
                    level = int(style_name.replace('Heading', '').strip())
                    lines.append(f"{'#' * level} {text}")
                    continue
                except ValueError:
                    pass
            
            # Check if text looks like a heading (short text with heading keywords)
            if len(text) < 20:
                is_heading = False
                for keyword in heading_keywords:
                    if keyword in text:
                        is_heading = True
                        break
                
                if is_heading:
                    lines.append(f"## {text}")
                    continue
            
            # Detect list items
            if style_name and style_name.startswith('List'):
                lines.append(f"- {text}")
            else:
                lines.append(text)
        
        # Read tables
        for table in doc.tables:
            lines.append('')  # Add blank line before table
            
            # Process each row
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                row_text = ' | '.join(cells)
                lines.append(f"| {row_text} |")
                
                # Add separator after header row
                if i == 0:
                    separator = ' | '.join(['---'] * len(cells))
                    lines.append(f"| {separator} |")
            
            lines.append('')  # Add blank line after table
        
        return "\n".join(lines)
    except ImportError:
        logger.error("python-docx is not installed. Run: pip install python-docx")
        return ""
    except Exception as exc:
        logger.error("Failed to read docx file: %s", exc)
        return ""


def _read_pdf(file_path: Path) -> str:
    """Read a .pdf file using PyPDF2. Returns plain text content."""
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text())
            return "\n".join(text_parts)
    except ImportError:
        logger.error("PyPDF2 is not installed. Run: pip install PyPDF2")
        return ""
    except Exception as exc:
        logger.error("Failed to read PDF file: %s", exc)
        return ""


def read_file_content(file_path: str) -> str:
    """Read content from various file formats.
    
    复用 file_reader.read_file() 函数，避免重复造轮子。
    
    Parameters
    ----------
    file_path : str
        Path to the file to read.
    
    Returns
    -------
    str
        File content as string, or empty string if failed.
    """
    from core.file_reader import read_file
    
    content = read_file(file_path)
    
    # read_file 返回错误信息时，返回空字符串
    if content.startswith("[Error]") or content.startswith("[Skipped]"):
        logger.warning("Failed to read file %s: %s", file_path, content)
        return ""
    
    return content


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _ensure_dirs() -> None:
    """Ensure template directories exist."""
    _BUILTIN_DIR.mkdir(parents=True, exist_ok=True)
    _CUSTOM_DIR.mkdir(parents=True, exist_ok=True)


def _load_template_file(path: Path) -> Optional[Dict]:
    """Load and validate a single template JSON file.

    Returns None if the file is invalid.
    """
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        if not isinstance(data, dict):
            logger.warning("Template file is not a dict: %s", path)
            return None
        if not _REQUIRED_KEYS.issubset(data.keys()):
            missing = _REQUIRED_KEYS - data.keys()
            logger.warning("Template file missing keys %s: %s", missing, path)
            return None
        return data
    except (json.JSONDecodeError, OSError) as exc:
        logger.warning("Failed to load template %s: %s", path, exc)
        return None


def _load_templates_from_dir(directory: Path) -> Dict[str, Dict]:
    """Load all valid templates from a directory.

    Returns a dict mapping template name (filename stem) to template data.
    """
    templates: Dict[str, Dict] = {}
    if not directory.exists():
        return templates

    for path in sorted(directory.glob("*.json")):
        data = _load_template_file(path)
        if data is not None:
            templates[path.stem] = data

    return templates


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def list_templates() -> Dict[str, Dict]:
    """List all available templates (builtin + custom).

    Returns
    -------
    dict
        Mapping of template_name -> template_data.
        Builtin templates come first, then custom templates.
        Custom templates with the same name override builtin ones.
    """
    _ensure_dirs()

    templates: Dict[str, Dict] = {}

    # Load builtin templates
    builtin = _load_templates_from_dir(_BUILTIN_DIR)
    templates.update(builtin)

    # Load custom templates (override builtin if same name)
    custom = _load_templates_from_dir(_CUSTOM_DIR)
    templates.update(custom)

    return templates


def get_template(name: str) -> Optional[Dict]:
    """Get a specific template by name.

    Parameters
    ----------
    name : str
        Template name (filename stem, e.g. "standard", "concise").

    Returns
    -------
    dict or None
        Template data if found, None otherwise.
    """
    _ensure_dirs()

    # Check custom first (higher priority)
    custom_path = _CUSTOM_DIR / f"{name}.json"
    if custom_path.exists():
        data = _load_template_file(custom_path)
        if data is not None:
            return data

    # Then check builtin
    builtin_path = _BUILTIN_DIR / f"{name}.json"
    if builtin_path.exists():
        data = _load_template_file(builtin_path)
        if data is not None:
            return data

    return None


def save_custom_template(name: str, template_data: Dict) -> bool:
    """Save a custom template.

    Parameters
    ----------
    name : str
        Template name (will be used as filename).
    template_data : dict
        Template data dict with required keys.

    Returns
    -------
    bool
        True if saved successfully, False otherwise.

    Raises
    ------
    ValueError
        If template_data is missing required keys.
    """
    _ensure_dirs()

    # Validate required keys
    missing = _REQUIRED_KEYS - template_data.keys()
    if missing:
        raise ValueError(f"Template data missing required keys: {missing}")

    # Sanitize filename
    safe_name = re.sub(r'[^\w\-]', '_', name).strip('_')
    if not safe_name:
        raise ValueError("Invalid template name")

    path = _CUSTOM_DIR / f"{safe_name}.json"

    try:
        path.write_text(
            json.dumps(template_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        logger.info("Saved custom template: %s", path)
        return True
    except OSError as exc:
        logger.error("Failed to save template %s: %s", path, exc)
        return False


def delete_custom_template(name: str) -> bool:
    """Delete a custom template.

    Parameters
    ----------
    name : str
        Template name (filename stem).

    Returns
    -------
    bool
        True if deleted successfully, False if not found or error.
    """
    _ensure_dirs()

    path = _CUSTOM_DIR / f"{name}.json"
    if not path.exists():
        logger.warning("Custom template not found: %s", name)
        return False

    try:
        path.unlink()
        logger.info("Deleted custom template: %s", name)
        return True
    except OSError as exc:
        logger.error("Failed to delete template %s: %s", path, exc)
        return False


def extract_template_from_report(report_text: str, name: str) -> Dict:
    """Extract a template structure from an existing weekly report.

    Analyzes the markdown structure of a report and creates a template
    by replacing specific content with placeholders while preserving structure.

    Parameters
    ----------
    report_text : str
        The markdown content of an existing weekly report.
    name : str
        Name for the new template.

    Returns
    -------
    dict
        Template data dict ready to be saved.
    """
    lines = report_text.strip().split('\n')
    sections: List[str] = []
    template_lines: List[str] = []
    
    # Common section keywords to detect
    section_keywords = [
        '工作总结', '工作汇总', '本周工作', '近期工作',
        '重点工作', '项目进展', '工作内容',
        '遇到的问题', '问题与解决', '问题', '风险',
        '后续计划', '下周计划', '计划', '下一步',
        '完成事项', '已完成', '完成',
        '进行中', '进行中的工作',
        '需要的支持', '需要协助',
        '技术决策', '决策',
    ]

    # Track if we're in a table
    in_table = False
    table_rows = []
    
    for line in lines:
        stripped = line.strip()

        # Skip empty lines but preserve them
        if not stripped:
            if in_table:
                # End of table
                template_lines.extend(table_rows)
                template_lines.append('')
                table_rows = []
                in_table = False
            template_lines.append('')
            continue

        # Handle table rows
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            
            # Check if it's a separator row
            if set(stripped.replace('|', '').replace('-', '').replace(':', '').strip()) == set():
                table_rows.append(stripped)
                continue
            
            # Count columns
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            num_cols = len(cells)
            
            # Replace content with placeholders but keep structure
            if num_cols == 1:
                # Single column table - might be header or description
                table_rows.append(f"| [内容] |")
            else:
                # Multi-column table
                placeholders = ' | '.join(['[内容]'] * num_cols)
                table_rows.append(f"| {placeholders} |")
            continue
        
        # If we were in a table and hit non-table line, flush table
        if in_table:
            template_lines.extend(table_rows)
            template_lines.append('')
            table_rows = []
            in_table = False

        # Detect headings - preserve structure
        if stripped.startswith('#'):
            # Count heading level
            level = 0
            for ch in stripped:
                if ch == '#':
                    level += 1
                else:
                    break

            heading_text = stripped[level:].strip()

            # For level 1 (title), replace date ranges with placeholder
            if level == 1:
                # Replace date patterns like MM.DD or YYYY-MM-DD
                heading_text = re.sub(
                    r'\d{4}[-./]\d{1,2}[-./]\d{1,2}',
                    '{week_range}',
                    heading_text,
                )
                heading_text = re.sub(
                    r'\d{1,2}\.\d{1,2}\s*[-–—]\s*\d{1,2}\.\d{1,2}',
                    '{week_range}',
                    heading_text,
                )
                template_lines.append(f"{'#' * level} {heading_text}")
                continue

            # For level 2 headings, record as sections
            if level == 2:
                sections.append(heading_text)
                template_lines.append(f"{'#' * level} {heading_text}")
                continue

            # For deeper levels, keep as is
            template_lines.append(f"{'#' * level} {heading_text}")
            continue
        
        # Detect section-like lines (Chinese keywords that look like headings)
        if not stripped.startswith('#') and not stripped.startswith('-') and not stripped.startswith('*'):
            # Check if this line looks like a section title
            is_section = False
            for keyword in section_keywords:
                if keyword in stripped and len(stripped) < 20:
                    is_section = True
                    sections.append(stripped)
                    template_lines.append(f"## {stripped}")
                    break
            
            if is_section:
                continue

        # Handle list items - preserve format with placeholder
        if stripped.startswith('- ') or stripped.startswith('* '):
            content = stripped[2:]
            prefix = stripped[0]  # - or *

            # Check if it's a checkbox item
            if content.startswith('[ ]') or content.startswith('[x]'):
                template_lines.append(f"{prefix} [ ] [任务描述]")
                continue

            # Check if it's a key-value style (问题：... → 方案：...)
            if '→' in content or '->' in content:
                template_lines.append(f"{prefix} [问题描述] → [解决方案]")
                continue

            # Check for specific patterns - preserve the key, replace value
            for key_prefix in ['进展：', '进展:', '成果：', '成果:', '问题：', '问题:', '方案：', '方案:']:
                if content.startswith(key_prefix):
                    template_lines.append(f"{prefix} {key_prefix}[描述]")
                    break
            else:
                # Generic list item - replace with placeholder
                template_lines.append(f"{prefix} [内容]")
            continue

        # Handle numbered list items
        if re.match(r'^\d+\.\s', stripped):
            template_lines.append(f"- [内容]")
            continue

        # Handle separator lines
        if stripped.startswith('---'):
            template_lines.append('---')
            continue

        # Handle timestamp lines
        if '*自动生成于' in stripped or '*自动生成' in stripped:
            template_lines.append('*自动生成于 YYYY-MM-DD HH:MM*')
            continue

        # Generic paragraph - preserve structure but replace specific content
        # Keep the line as is if it looks like a template instruction
        if any(keyword in stripped for keyword in ['请', '用', '列出', '根据', '如果', '没有']):
            template_lines.append(stripped)
        else:
            # For other paragraphs, keep the structure but mark as placeholder
            template_lines.append('[内容描述]')

    # Flush any remaining table
    if in_table and table_rows:
        template_lines.extend(table_rows)
        template_lines.append('')

    # Build the prompt template
    template_content = '\n'.join(template_lines)
    
    prompt_template = f"""你是一个工作汇总助手。请根据以下文件分析结果，生成一份结构化的工作汇总。

## 汇总日期范围
{{week_range}}

## 文件分析结果
{{analyses}}

## 汇总生成要求

请严格按照以下格式生成工作汇总：

{template_content}

## 生成规则

1. **归类整理**：将相关文件的分析结果归类到同一模块，不要按文件逐个罗列
2. **突出重点**：重点写成果和进展，不要写流水账
3. **语言专业**：用专业的职场语言描述，适当包装工作成果
4. **格式规范**：严格按照上述Markdown格式输出，确保可解析
5. **处理空内容**：如果某个部分没有相关内容，使用默认描述（如"暂无"）

## 注意事项
- 不要编造没有依据的内容
- 不要输出代码片段
- 保持客观真实，基于提供的分析结果生成"""

    return {
        "name": name,
        "description": f"从现有周报提取的模板 - {name}",
        "template": prompt_template,
        "sections": sections,
        "created_at": datetime.now().strftime("%Y-%m-%d"),
    }


def analyze_template_with_llm(
    template_content: str,
    provider,
    template_name: str
) -> Dict:
    """使用 LLM 分析模板格式并生成优化的模板结构。

    Parameters
    ----------
    template_content : str
        原始模板内容（Markdown 或纯文本）。
    provider : LLMProvider
        LLM 提供者实例。
    template_name : str
        模板名称。

    Returns
    -------
    dict
        分析后的模板数据，包含：
        - name: 模板名称
        - description: 模板描述
        - template: 优化后的提示词模板
        - sections: 检测到的章节列表
        - created_at: 创建时间
        - analysis: LLM 分析结果
    """
    # 构建分析提示词
    analysis_prompt = f"""你是一个周报模板分析专家。请分析以下模板内容，提取其结构和格式要求。

## 模板内容
{template_content}

## 分析要求

请按以下格式输出分析结果：

### 模板概述
简要描述这个模板的用途和特点。

### 章节结构
列出模板中的主要章节（用逗号分隔）。

### 格式要求
描述模板的格式要求，包括：
1. 标题层级
2. 列表格式
3. 表格结构（如有）
4. 特殊标记或符号

### 优化建议
提供模板优化建议，使其更适合AI生成工作汇总。

### 提示词模板
基于分析结果，生成一个优化的提示词模板，包含以下占位符：
- {{week_range}}: 日期范围
- {{analyses}}: 文件分析结果

请确保提示词模板能够引导AI按照原模板的结构生成工作汇总。

## 注意事项
- 保持原模板的核心结构
- 优化语言使其更适合AI理解
- 确保生成的内容专业、简洁
"""

    try:
        # 调用 LLM 进行分析
        messages = [{"role": "user", "content": analysis_prompt}]
        response = provider.chat_completion(
            messages=messages,
            temperature=0.3,
            max_tokens=2048,
        )

        if not response or not response.strip():
            logger.warning("LLM 返回空响应")
            return _fallback_template_extraction(template_content, template_name)

        analysis_result = response.strip()

        # 从分析结果中提取章节
        sections = _extract_sections_from_analysis(analysis_result)

        # 从分析结果中提取提示词模板
        prompt_template = _extract_prompt_template(analysis_result, template_content)

        return {
            "name": template_name,
            "description": f"LLM 分析优化的模板 - {template_name}",
            "template": prompt_template,
            "sections": sections,
            "created_at": datetime.now().strftime("%Y-%m-%d"),
            "analysis": analysis_result,
        }

    except Exception as e:
        logger.error("LLM 模板分析失败: %s", e)
        return _fallback_template_extraction(template_content, template_name)


def _extract_sections_from_analysis(analysis_text: str) -> List[str]:
    """从 LLM 分析结果中提取章节列表。"""
    sections = []
    in_sections = False

    for line in analysis_text.split('\n'):
        stripped = line.strip()

        # 查找章节结构部分
        if '章节结构' in stripped or '章節結構' in stripped:
            in_sections = True
            continue

        if in_sections:
            # 遇到下一个标题时停止
            if stripped.startswith('###') or stripped.startswith('##'):
                break

            # 提取章节名称
            if stripped and not stripped.startswith('#'):
                # 移除序号和标点
                import re
                cleaned = re.sub(r'^[\d\.\-\*\s]+', '', stripped)
                cleaned = cleaned.strip('，。、；：')
                if cleaned:
                    sections.extend([s.strip() for s in cleaned.split('，') if s.strip()])

    return sections


def _extract_prompt_template(analysis_text: str, original_content: str) -> str:
    """从 LLM 分析结果中提取提示词模板。"""
    import re

    # 尝试查找 "提示词模板" 部分
    in_template = False
    template_lines = []

    for line in analysis_text.split('\n'):
        stripped = line.strip()

        if '提示词模板' in stripped or '提示詞模板' in stripped:
            in_template = True
            continue

        if in_template:
            # 遇到下一个标题时停止
            if stripped.startswith('###') or stripped.startswith('##'):
                break

            template_lines.append(line)

    if template_lines:
        extracted = '\n'.join(template_lines).strip()
        # 确保包含必要的占位符
        if '{week_range}' not in extracted:
            extracted = f"## 汇总日期范围\n{{week_range}}\n\n{extracted}"
        if '{analyses}' not in extracted:
            extracted = f"{extracted}\n\n## 文件分析结果\n{{analyses}}"
        return extracted

    # 如果无法提取，使用原始内容构建
    return _build_template_from_content(original_content)


def _build_template_from_content(content: str) -> str:
    """从原始内容构建提示词模板。"""
    return f"""你是一个工作汇总助手。请根据以下文件分析结果，生成一份结构化的工作汇总。

## 汇总日期范围
{{week_range}}

## 文件分析结果
{{analyses}}

## 模板格式要求

请严格按照以下格式生成工作汇总：

{content}

## 生成规则

1. **归类整理**：将相关文件的分析结果归类到同一模块，不要按文件逐个罗列
2. **突出重点**：重点写成果和进展，不要写流水账
3. **语言专业**：用专业的职场语言描述，适当包装工作成果
4. **格式规范**：严格按照上述Markdown格式输出，确保可解析
5. **处理空内容**：如果某个部分没有相关内容，使用默认描述（如"暂无"）

## 注意事项
- 不要编造没有依据的内容
- 不要输出代码片段
- 保持客观真实，基于提供的分析结果生成"""


def _fallback_template_extraction(content: str, name: str) -> Dict:
    """当 LLM 分析失败时的备用模板提取方法。"""
    # 使用原有的 extract_template_from_report 逻辑
    return extract_template_from_report(content, name)
