"""周报终结者 V1 - Streamlit Web界面."""

import io
import json
import logging
import queue
import re
import sys
from datetime import datetime, time
from pathlib import Path
from typing import Dict, List, Tuple

# 抑制 Streamlit 的 ScriptRunContext 警告
logging.getLogger("streamlit.runtime.scriptrunner_utils.script_run_context").setLevel(logging.ERROR)

import streamlit as st

logger = logging.getLogger(__name__)

# Ensure project root is on sys.path so local imports work
_project_root = str(Path(__file__).resolve().parent)
if _project_root not in sys.path:
    sys.path.insert(0, _project_root)

from config import Config  # noqa: E402
from core.file_reader import scan_folder, scan_folder_all, get_file_hash, read_file  # noqa: E402
from core.analyzer import analyze_all_files, analyze_changes  # noqa: E402
from core.diff_extractor import get_file_diff, compare_file_contents  # noqa: E402
from core.generator import generate_report, get_week_range  # noqa: E402
from core.history import (  # noqa: E402
    save_snapshot,
    get_latest_snapshot,
    list_snapshots,
    compare_with_latest,
    delete_snapshot,
    clear_old_snapshots,
    get_snapshot_file_contents,
    get_timeline_summary,
)
from core.template_manager import (  # noqa: E402
    list_templates,
    get_template,
    save_custom_template,
    delete_custom_template,
    extract_template_from_report,
)
from llm.factory import (  # noqa: E402
    create_provider,
    get_configured_providers,
    PROVIDER_DISPLAY_NAMES,
)
from llm.custom_provider import CustomProvider  # noqa: E402
from core.scheduler import AutoScheduler  # noqa: E402
from core.report_archive import (  # noqa: E402
    archive_report,
    list_archived_reports,
    get_archived_report,
    delete_archived_report,
    get_archive_stats,
)


# ---------------------------------------------------------------------------
# Page config
# ---------------------------------------------------------------------------

st.set_page_config(
    page_title="周报生成器",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ---------------------------------------------------------------------------
# Design System - 全局CSS注入
# ---------------------------------------------------------------------------

DESIGN_SYSTEM_CSS = """
<style>
    /* ====== Design Tokens ====== */
    :root {
        /* Primary palette - Deep Indigo */
        --primary-50: #EEF2FF;
        --primary-100: #E0E7FF;
        --primary-200: #C7D2FE;
        --primary-300: #A5B4FC;
        --primary-400: #818CF8;
        --primary-500: #6366F1;
        --primary-600: #4F46E5;
        --primary-700: #4338CA;
        --primary-800: #3730A3;
        --primary-900: #312E81;

        /* Accent - Warm Gold */
        --accent-400: #FBBF24;
        --accent-500: #F59E0B;
        --accent-600: #D97706;

        /* Neutral palette */
        --neutral-50: #F8FAFC;
        --neutral-100: #F1F5F9;
        --neutral-200: #E2E8F0;
        --neutral-300: #CBD5E1;
        --neutral-400: #94A3B8;
        --neutral-500: #64748B;
        --neutral-600: #475569;
        --neutral-700: #334155;
        --neutral-800: #1E293B;
        --neutral-900: #0F172A;

        /* Semantic */
        --success: #10B981;
        --warning: #F59E0B;
        --error: #EF4444;
        --info: #3B82F6;

        /* ====== iOS Glass Material Tokens ====== */
        --glass-blur: 20px;
        --glass-bg: rgba(255, 255, 255, 0.65);
        --glass-bg-heavy: rgba(255, 255, 255, 0.8);
        --glass-bg-dark: rgba(15, 23, 42, 0.75);
        --glass-border: rgba(255, 255, 255, 0.35);
        --glass-border-subtle: rgba(255, 255, 255, 0.18);
        --glass-inset: inset 0 1px 0 rgba(255, 255, 255, 0.6);
        --glass-inset-subtle: inset 0 1px 0 rgba(255, 255, 255, 0.25);
        --glow-primary: rgba(99, 102, 241, 0.4);
        --glow-primary-soft: rgba(99, 102, 241, 0.2);
        --glow-accent: rgba(251, 191, 36, 0.35);
        --glow-success: rgba(16, 185, 129, 0.35);

        /* Shadows - Multi-layer depth system */
        --shadow-sm: 0 1px 3px rgba(0, 0, 0, 0.06), 0 1px 2px rgba(0, 0, 0, 0.04);
        --shadow-md: 0 4px 12px rgba(0, 0, 0, 0.08), 0 2px 4px rgba(0, 0, 0, 0.04);
        --shadow-lg: 0 10px 25px rgba(0, 0, 0, 0.1), 0 4px 10px rgba(0, 0, 0, 0.05);
        --shadow-xl: 0 20px 40px rgba(0, 0, 0, 0.12), 0 8px 16px rgba(0, 0, 0, 0.06);
        --shadow-glass: 0 8px 32px rgba(0, 0, 0, 0.1), 0 2px 8px rgba(0, 0, 0, 0.05), inset 0 1px 0 rgba(255, 255, 255, 0.6);
        --shadow-glass-hover: 0 12px 40px rgba(0, 0, 0, 0.14), 0 4px 12px rgba(0, 0, 0, 0.06), inset 0 1px 0 rgba(255, 255, 255, 0.7);
        --shadow-glow-btn: 0 0 15px var(--glow-primary), 0 0 30px var(--glow-primary-soft), 0 10px 40px rgba(0, 0, 0, 0.15);
        --shadow-glow-btn-hover: 0 0 20px rgba(99, 102, 241, 0.6), 0 0 40px var(--glow-primary), 0 15px 50px rgba(0, 0, 0, 0.2);

        /* Border radius */
        --radius-sm: 8px;
        --radius-md: 12px;
        --radius-lg: 18px;
        --radius-xl: 24px;

        /* Transitions */
        --transition-fast: 180ms cubic-bezier(0.4, 0, 0.2, 1);
        --transition-normal: 300ms cubic-bezier(0.4, 0, 0.2, 1);
        --transition-slow: 400ms cubic-bezier(0.4, 0, 0.2, 1);
    }

    /* ====== Global Overrides ====== */
    .stApp {
        background: linear-gradient(145deg, #E8ECF4 0%, #DDE4F0 25%, #EAE8F5 50%, #E2EAF2 75%, #D8E2EE 100%) !important;
        min-height: 100vh;
    }

    /* Animated background mesh - DISABLED to prevent flickering
    .stApp::before {
        content: '';
        position: fixed;
        top: -30%;
        right: -20%;
        width: 700px;
        height: 700px;
        background: radial-gradient(circle, rgba(99, 102, 241, 0.08) 0%, transparent 60%);
        pointer-events: none;
        z-index: 0;
        animation: floatOrb1 25s ease-in-out infinite;
    }
    .stApp::after {
        content: '';
        position: fixed;
        bottom: -20%;
        left: -15%;
        width: 600px;
        height: 600px;
        background: radial-gradient(circle, rgba(251, 191, 36, 0.06) 0%, transparent 60%);
        pointer-events: none;
        z-index: 0;
        animation: floatOrb2 30s ease-in-out infinite;
    }
    */
    @keyframes floatOrb1 {
        0%, 100% { transform: translate(0, 0) scale(1); }
        33% { transform: translate(-40px, 30px) scale(1.05); }
        66% { transform: translate(20px, -20px) scale(0.95); }
    }
    @keyframes floatOrb2 {
        0%, 100% { transform: translate(0, 0) scale(1); }
        50% { transform: translate(30px, -40px) scale(1.08); }
    }

    /* Main container tighter padding */
    .block-container {
        padding-top: 2rem !important;
        padding-bottom: 2rem !important;
        max-width: 900px !important;
        position: relative;
        z-index: 1;
    }

    /* ====== Hero Banner - iOS Glass with Depth ====== */
    .hero-banner {
        background: linear-gradient(135deg, #1E1B4B 0%, #312E81 30%, #4338CA 60%, #6366F1 100%);
        border-radius: var(--radius-xl);
        padding: 2.75rem 2.5rem 2.25rem;
        margin-bottom: 1.75rem;
        position: relative;
        overflow: hidden;
        box-shadow:
            0 25px 60px rgba(49, 46, 129, 0.35),
            0 10px 25px rgba(0, 0, 0, 0.15),
            inset 0 1px 0 rgba(255, 255, 255, 0.15),
            inset 0 -1px 0 rgba(0, 0, 0, 0.1);
        border: 1px solid rgba(255, 255, 255, 0.1);
    }
    .hero-banner::before {
        content: '';
        position: absolute;
        top: -60%;
        right: -25%;
        width: 500px;
        height: 500px;
        background: radial-gradient(circle, rgba(251, 191, 36, 0.18) 0%, rgba(251, 191, 36, 0.05) 40%, transparent 70%);
        pointer-events: none;
    }
    .hero-banner::after {
        content: '';
        position: absolute;
        bottom: -40%;
        left: -15%;
        width: 400px;
        height: 400px;
        background: radial-gradient(circle, rgba(129, 140, 248, 0.25) 0%, rgba(99, 102, 241, 0.08) 40%, transparent 70%);
        pointer-events: none;
    }
    @keyframes heroGlow1 {
        0%, 100% { transform: translate(0, 0); opacity: 1; }
        50% { transform: translate(-20px, 15px); opacity: 0.7; }
    }
    @keyframes heroGlow2 {
        0%, 100% { transform: translate(0, 0); opacity: 0.8; }
        50% { transform: translate(15px, -10px); opacity: 1; }
    }
    /* Hero noise texture overlay */
    .hero-banner > *:first-child::before {
        content: '';
        position: absolute;
        inset: 0;
        background-image: url("data:image/svg+xml,%3Csvg viewBox='0 0 256 256' xmlns='http://www.w3.org/2000/svg'%3E%3Cfilter id='n'%3E%3CfeTurbulence type='fractalNoise' baseFrequency='0.9' numOctaves='4' stitchTiles='stitch'/%3E%3C/filter%3E%3Crect width='100%25' height='100%25' filter='url(%23n)' opacity='0.03'/%3E%3C/svg%3E");
        pointer-events: none;
        z-index: 0;
    }
    .hero-banner h1 {
        color: #FFFFFF !important;
        font-size: 2rem !important;
        font-weight: 700 !important;
        margin: 0 0 0.5rem 0 !important;
        letter-spacing: -0.02em;
        position: relative;
        z-index: 1;
        text-shadow: 0 2px 10px rgba(0, 0, 0, 0.2), 0 1px 3px rgba(0, 0, 0, 0.15);
    }
    .hero-banner p {
        color: rgba(255, 255, 255, 0.85) !important;
        font-size: 1rem !important;
        margin: 0 !important;
        position: relative;
        z-index: 1;
        text-shadow: 0 1px 4px rgba(0, 0, 0, 0.15);
    }
    .hero-badge {
        display: inline-block;
        background: rgba(255, 255, 255, 0.12);
        backdrop-filter: blur(12px);
        -webkit-backdrop-filter: blur(12px);
        border: 1px solid rgba(255, 255, 255, 0.2);
        border-radius: 999px;
        padding: 0.3rem 0.85rem;
        font-size: 0.75rem;
        font-weight: 600;
        color: rgba(255, 255, 255, 0.95);
        margin-bottom: 0.85rem;
        position: relative;
        z-index: 1;
        box-shadow:
            0 0 15px rgba(251, 191, 36, 0.25),
            0 0 30px rgba(251, 191, 36, 0.1),
            inset 0 1px 0 rgba(255, 255, 255, 0.2);
        letter-spacing: 0.03em;
    }

    /* ====== Card System - Frosted Glass ====== */
    .card {
        background: rgba(255, 255, 255, 0.65);
        backdrop-filter: blur(var(--glass-blur));
        -webkit-backdrop-filter: blur(var(--glass-blur));
        border-radius: var(--radius-lg);
        border: 1px solid var(--glass-border);
        padding: 1.75rem;
        margin-bottom: 1.25rem;
        box-shadow: var(--shadow-glass);
        transition: all var(--transition-normal);
        position: relative;
    }
    .card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 1px;
        background: linear-gradient(90deg, transparent 10%, rgba(255, 255, 255, 0.8) 50%, transparent 90%);
        border-radius: var(--radius-lg) var(--radius-lg) 0 0;
        pointer-events: none;
    }
    .card:hover {
        box-shadow: var(--shadow-glass-hover);
        transform: translateY(-2px);
        border-color: rgba(255, 255, 255, 0.5);
        background: rgba(255, 255, 255, 0.75);
    }
    .card-header {
        display: flex;
        align-items: center;
        gap: 0.6rem;
        margin-bottom: 1.25rem;
        padding-bottom: 0.85rem;
        border-bottom: 1px solid rgba(148, 163, 184, 0.12);
    }
    .card-header-icon {
        font-size: 1.25rem;
        width: 36px;
        height: 36px;
        display: flex;
        align-items: center;
        justify-content: center;
        background: linear-gradient(135deg, rgba(99, 102, 241, 0.1), rgba(129, 140, 248, 0.05));
        border-radius: var(--radius-sm);
        border: 1px solid rgba(99, 102, 241, 0.12);
        box-shadow: 0 0 12px rgba(99, 102, 241, 0.08);
    }
    .card-header-title {
        font-size: 1rem;
        font-weight: 600;
        color: var(--neutral-800);
        letter-spacing: -0.01em;
    }

    /* ====== Stat Cards - Glass with Gradient Overlays ====== */
    .stat-grid {
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        gap: 1rem;
        margin-bottom: 0.5rem;
    }
    .stat-card {
        background: rgba(255, 255, 255, 0.55);
        backdrop-filter: blur(16px);
        -webkit-backdrop-filter: blur(16px);
        border: 1px solid rgba(255, 255, 255, 0.4);
        border-radius: var(--radius-md);
        padding: 1.35rem;
        text-align: center;
        transition: all var(--transition-normal);
        position: relative;
        overflow: hidden;
        box-shadow: 0 4px 16px rgba(0, 0, 0, 0.06), inset 0 1px 0 rgba(255, 255, 255, 0.7);
    }
    .stat-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 3px;
        background: linear-gradient(90deg, var(--primary-400), var(--accent-400));
        opacity: 0;
        transition: opacity var(--transition-normal);
    }
    .stat-card::after {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 1px;
        background: linear-gradient(90deg, transparent, rgba(255, 255, 255, 0.9), transparent);
        pointer-events: none;
    }
    .stat-card:hover {
        transform: translateY(-4px);
        box-shadow: 0 8px 28px rgba(0, 0, 0, 0.1), 0 0 20px rgba(99, 102, 241, 0.08), inset 0 1px 0 rgba(255, 255, 255, 0.8);
        border-color: rgba(99, 102, 241, 0.2);
        background: rgba(255, 255, 255, 0.7);
    }
    .stat-card:hover::before {
        opacity: 1;
    }
    .stat-icon {
        font-size: 1.5rem;
        margin-bottom: 0.5rem;
        display: inline-flex;
        align-items: center;
        justify-content: center;
        width: 44px;
        height: 44px;
        border-radius: 50%;
        background: linear-gradient(135deg, rgba(99, 102, 241, 0.12), rgba(99, 102, 241, 0.04));
        box-shadow: 0 0 18px rgba(99, 102, 241, 0.1);
        border: 1px solid rgba(99, 102, 241, 0.1);
    }
    .stat-value {
        font-size: 1.75rem;
        font-weight: 700;
        color: var(--neutral-900);
        line-height: 1.2;
        letter-spacing: -0.02em;
        text-shadow: 0 1px 2px rgba(0, 0, 0, 0.04);
    }
    .stat-label {
        font-size: 0.78rem;
        color: var(--neutral-500);
        margin-top: 0.3rem;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.06em;
    }

    /* ====== Config Info Card ====== */
    .config-info {
        display: flex;
        align-items: center;
        gap: 1.5rem;
        flex-wrap: wrap;
    }
    .config-item {
        display: flex;
        align-items: center;
        gap: 0.6rem;
    }
    .config-item-icon {
        width: 36px;
        height: 36px;
        border-radius: var(--radius-sm);
        background: linear-gradient(135deg, rgba(99, 102, 241, 0.12), rgba(99, 102, 241, 0.04));
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 1rem;
        border: 1px solid rgba(99, 102, 241, 0.1);
        box-shadow: 0 0 12px rgba(99, 102, 241, 0.06);
    }
    .config-item-label {
        font-size: 0.72rem;
        color: var(--neutral-500);
        text-transform: uppercase;
        letter-spacing: 0.06em;
        font-weight: 600;
    }
    .config-item-value {
        font-size: 0.9rem;
        font-weight: 600;
        color: var(--neutral-800);
    }

    /* ====== Buttons - Glow + Glass ====== */
    .stButton > button[kind="primary"],
    .stButton > button[data-testid="stBaseButton-primary"] {
        background: linear-gradient(135deg, #4338CA 0%, #6366F1 50%, #818CF8 100%) !important;
        border: none !important;
        border-radius: var(--radius-md) !important;
        padding: 0.8rem 1.5rem !important;
        font-weight: 600 !important;
        font-size: 0.95rem !important;
        letter-spacing: 0.01em !important;
        box-shadow: var(--shadow-glow-btn) !important;
        transition: all var(--transition-normal) !important;
        color: #FFFFFF !important;
        position: relative;
        overflow: hidden;
        text-shadow: 0 1px 2px rgba(0, 0, 0, 0.15);
    }
    .stButton > button[kind="primary"]::before,
    .stButton > button[data-testid="stBaseButton-primary"]::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 50%;
        background: linear-gradient(180deg, rgba(255, 255, 255, 0.15), transparent);
        pointer-events: none;
        border-radius: var(--radius-md) var(--radius-md) 0 0;
    }
    .stButton > button[kind="primary"]:hover,
    .stButton > button[data-testid="stBaseButton-primary"]:hover {
        transform: translateY(-3px) !important;
        box-shadow: var(--shadow-glow-btn-hover) !important;
        background: linear-gradient(135deg, #3730A3 0%, #4F46E5 50%, #6366F1 100%) !important;
    }
    .stButton > button[kind="primary"]:active,
    .stButton > button[data-testid="stBaseButton-primary"]:active {
        transform: translateY(-1px) !important;
        box-shadow: 0 0 12px var(--glow-primary), 0 5px 20px rgba(0, 0, 0, 0.15) !important;
    }

    /* Secondary buttons - Glass outline */
    .stButton > button[kind="secondary"],
    .stButton > button:not([kind]) {
        border-radius: var(--radius-md) !important;
        border: 1px solid rgba(148, 163, 184, 0.25) !important;
        font-weight: 500 !important;
        transition: all var(--transition-normal) !important;
        background: rgba(255, 255, 255, 0.5) !important;
        backdrop-filter: blur(8px) !important;
        -webkit-backdrop-filter: blur(8px) !important;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.04), inset 0 1px 0 rgba(255, 255, 255, 0.6) !important;
    }
    .stButton > button[kind="secondary"]:hover,
    .stButton > button:not([kind]):hover {
        border-color: rgba(99, 102, 241, 0.35) !important;
        color: var(--primary-700) !important;
        background: rgba(238, 242, 255, 0.7) !important;
        box-shadow: 0 4px 16px rgba(99, 102, 241, 0.1), 0 0 20px rgba(99, 102, 241, 0.06), inset 0 1px 0 rgba(255, 255, 255, 0.7) !important;
        transform: translateY(-1px) !important;
    }

    /* ====== Inputs - Glass with Focus Glow ====== */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea {
        border-radius: var(--radius-md) !important;
        border: 1.5px solid rgba(148, 163, 184, 0.25) !important;
        padding: 0.8rem 1rem !important;
        font-size: 0.9rem !important;
        transition: all var(--transition-normal) !important;
        background: rgba(255, 255, 255, 0.55) !important;
        backdrop-filter: blur(10px) !important;
        -webkit-backdrop-filter: blur(10px) !important;
        box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.5), 0 2px 6px rgba(0, 0, 0, 0.03) !important;
    }
    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {
        border-color: rgba(99, 102, 241, 0.5) !important;
        box-shadow:
            0 0 0 3px rgba(99, 102, 241, 0.12),
            0 0 20px rgba(99, 102, 241, 0.08),
            inset 0 1px 0 rgba(255, 255, 255, 0.6) !important;
        background: rgba(255, 255, 255, 0.7) !important;
    }
    .stTextInput > div > div > input::placeholder,
    .stTextArea > div > div > textarea::placeholder {
        color: var(--neutral-400) !important;
    }

    /* ====== Selectbox - Glass ====== */
    .stSelectbox > div > div {
        border-radius: var(--radius-md) !important;
        border: 1.5px solid rgba(148, 163, 184, 0.25) !important;
        background: rgba(255, 255, 255, 0.55) !important;
        backdrop-filter: blur(10px) !important;
        -webkit-backdrop-filter: blur(10px) !important;
    }

    /* ====== Slider ====== */
    .stSlider > div > div > div > div {
        background: linear-gradient(90deg, var(--primary-500), var(--primary-400)) !important;
        box-shadow: 0 0 10px rgba(99, 102, 241, 0.3) !important;
    }

    /* ====== Tabs - Frosted Segments ====== */
    .stTabs [data-baseweb="tab-list"] {
        gap: 0 !important;
        background: rgba(241, 245, 249, 0.6) !important;
        backdrop-filter: blur(12px) !important;
        -webkit-backdrop-filter: blur(12px) !important;
        border-radius: var(--radius-md) !important;
        padding: 4px !important;
        border: 1px solid rgba(255, 255, 255, 0.4) !important;
        box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.5), 0 2px 6px rgba(0, 0, 0, 0.04) !important;
    }
    .stTabs [data-baseweb="tab"] {
        border-radius: var(--radius-sm) !important;
        padding: 0.6rem 1.25rem !important;
        font-weight: 500 !important;
        font-size: 0.875rem !important;
        color: var(--neutral-600) !important;
        transition: all var(--transition-fast) !important;
    }
    .stTabs [aria-selected="true"] {
        background: rgba(255, 255, 255, 0.85) !important;
        color: var(--primary-700) !important;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.06), inset 0 1px 0 rgba(255, 255, 255, 0.8) !important;
        font-weight: 600 !important;
    }
    .stTabs [data-baseweb="tab-border"] {
        display: none !important;
    }
    .stTabs [data-baseweb="tab-highlight"] {
        display: none !important;
    }

    /* ====== Dividers ====== */
    hr {
        border: none !important;
        height: 1px !important;
        background: linear-gradient(90deg, transparent, rgba(148, 163, 184, 0.25), transparent) !important;
        margin: 1.5rem 0 !important;
    }

    /* ====== Expander - Glass ====== */
    .streamlit-expanderHeader {
        font-weight: 600 !important;
        font-size: 0.9rem !important;
        color: var(--neutral-700) !important;
        border-radius: var(--radius-md) !important;
        background: rgba(255, 255, 255, 0.4) !important;
        backdrop-filter: blur(8px) !important;
        -webkit-backdrop-filter: blur(8px) !important;
        border: 1px solid rgba(255, 255, 255, 0.3) !important;
        padding: 0.75rem 1rem !important;
    }

    /* ====== Alerts - Glass ====== */
    .stAlert {
        border-radius: var(--radius-md) !important;
        border-left-width: 4px !important;
        backdrop-filter: blur(10px) !important;
        -webkit-backdrop-filter: blur(10px) !important;
        background: rgba(255, 255, 255, 0.6) !important;
    }

    /* ====== Metric Override ====== */
    [data-testid="stMetric"] {
        background: transparent !important;
        border: none !important;
        padding: 0 !important;
    }
    [data-testid="stMetricValue"] {
        font-size: 1.5rem !important;
        font-weight: 700 !important;
        color: var(--neutral-900) !important;
    }
    [data-testid="stMetricLabel"] {
        font-size: 0.8rem !important;
        color: var(--neutral-500) !important;
        text-transform: uppercase !important;
        letter-spacing: 0.05em !important;
    }

    /* ====== Sidebar - Dark Glass ====== */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, rgba(15, 23, 42, 0.92) 0%, rgba(30, 41, 59, 0.95) 100%) !important;
        backdrop-filter: blur(25px) !important;
        -webkit-backdrop-filter: blur(25px) !important;
        border-right: 1px solid rgba(255, 255, 255, 0.06) !important;
        box-shadow: 4px 0 24px rgba(0, 0, 0, 0.15) !important;
    }
    [data-testid="stSidebar"]::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: linear-gradient(180deg, rgba(99, 102, 241, 0.03) 0%, transparent 30%, rgba(251, 191, 36, 0.02) 100%);
        pointer-events: none;
    }
    [data-testid="stSidebar"] .stMarkdown h2,
    [data-testid="stSidebar"] .stMarkdown h3 {
        color: #FFFFFF !important;
    }
    [data-testid="stSidebar"] .stMarkdown p,
    [data-testid="stSidebar"] .stMarkdown li {
        color: rgba(255, 255, 255, 0.7) !important;
    }
    [data-testid="stSidebar"] .stButton > button {
        background: rgba(255, 255, 255, 0.07) !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
        color: rgba(255, 255, 255, 0.9) !important;
        border-radius: var(--radius-md) !important;
        transition: all var(--transition-normal) !important;
        backdrop-filter: blur(8px) !important;
        -webkit-backdrop-filter: blur(8px) !important;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.1), inset 0 1px 0 rgba(255, 255, 255, 0.06) !important;
    }
    [data-testid="stSidebar"] .stButton > button:hover {
        background: rgba(255, 255, 255, 0.14) !important;
        border-color: rgba(255, 255, 255, 0.22) !important;
        box-shadow:
            0 0 15px rgba(99, 102, 241, 0.15),
            0 4px 12px rgba(0, 0, 0, 0.15),
            inset 0 1px 0 rgba(255, 255, 255, 0.1) !important;
        transform: translateX(2px) !important;
    }
    [data-testid="stSidebar"] hr {
        background: linear-gradient(90deg, transparent, rgba(255, 255, 255, 0.1), transparent) !important;
    }

    /* ====== Sidebar Nav Active State ====== */
    .sidebar-brand {
        text-align: center;
        padding: 1.25rem 0 0.75rem;
        margin-bottom: 0.75rem;
        position: relative;
    }
    .sidebar-brand-icon {
        font-size: 2.25rem;
        display: block;
        margin-bottom: 0.35rem;
        filter: drop-shadow(0 0 12px rgba(99, 102, 241, 0.3));
    }
    .sidebar-brand-text {
        font-size: 0.88rem;
        font-weight: 700;
        color: rgba(255, 255, 255, 0.95);
        letter-spacing: 0.03em;
    }
    .sidebar-brand-sub {
        font-size: 0.68rem;
        color: rgba(255, 255, 255, 0.4);
        margin-top: 0.2rem;
        letter-spacing: 0.04em;
        text-transform: uppercase;
    }
    .sidebar-section-label {
        font-size: 0.62rem;
        text-transform: uppercase;
        letter-spacing: 0.12em;
        color: rgba(255, 255, 255, 0.3);
        padding: 0.5rem 0 0.3rem;
        font-weight: 700;
    }

    /* ====== File List Item - Glass ====== */
    .file-item {
        display: flex;
        align-items: center;
        gap: 0.75rem;
        padding: 0.65rem 0;
        border-bottom: 1px solid rgba(148, 163, 184, 0.1);
        transition: all var(--transition-fast);
        border-radius: var(--radius-sm);
    }
    .file-item:last-child {
        border-bottom: none;
    }
    .file-item:hover {
        background: rgba(255, 255, 255, 0.5);
        margin: 0 -0.5rem;
        padding-left: 0.5rem;
        padding-right: 0.5rem;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.04);
    }
    .file-item-icon {
        font-size: 1.25rem;
        width: 38px;
        height: 38px;
        display: flex;
        align-items: center;
        justify-content: center;
        background: linear-gradient(135deg, rgba(99, 102, 241, 0.08), rgba(99, 102, 241, 0.03));
        border-radius: var(--radius-sm);
        flex-shrink: 0;
        border: 1px solid rgba(99, 102, 241, 0.08);
    }
    .file-item-name {
        font-weight: 600;
        font-size: 0.875rem;
        color: var(--neutral-800);
    }
    .file-item-meta {
        font-size: 0.73rem;
        color: var(--neutral-500);
    }

    /* ====== CTA Button Enhanced ====== */
    .cta-section {
        text-align: center;
        padding: 0.75rem 0;
    }

    /* ====== Export Row ====== */
    .export-row {
        display: flex;
        align-items: center;
        gap: 1rem;
    }

    /* ====== Tip Box - Frosted ====== */
    .tip-box {
        background: rgba(255, 255, 255, 0.35);
        backdrop-filter: blur(12px);
        -webkit-backdrop-filter: blur(12px);
        border: 1px solid rgba(255, 255, 255, 0.25);
        border-radius: var(--radius-md);
        padding: 1rem 1.25rem;
        margin-top: 1rem;
        box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.3);
    }
    [data-testid="stSidebar"] .tip-box {
        background: rgba(255, 255, 255, 0.06) !important;
        border: 1px solid rgba(255, 255, 255, 0.08) !important;
        box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.04) !important;
    }
    .tip-box-title {
        font-size: 0.78rem;
        font-weight: 700;
        color: var(--primary-800);
        margin-bottom: 0.5rem;
        text-transform: uppercase;
        letter-spacing: 0.06em;
    }
    .tip-box-content {
        font-size: 0.84rem;
        color: var(--neutral-700);
        line-height: 1.65;
    }
    [data-testid="stSidebar"] .tip-box-title {
        color: rgba(255, 255, 255, 0.85) !important;
    }
    [data-testid="stSidebar"] .tip-box-content {
        color: rgba(255, 255, 255, 0.6) !important;
    }

    /* ====== Empty State ====== */
    .empty-state {
        text-align: center;
        padding: 2rem 1rem;
        color: var(--neutral-400);
    }
    .empty-state-icon {
        font-size: 3rem;
        margin-bottom: 0.75rem;
        opacity: 0.5;
        filter: drop-shadow(0 2px 4px rgba(0, 0, 0, 0.06));
    }
    .empty-state-text {
        font-size: 0.9rem;
        font-weight: 500;
    }

    /* ====== Scrollbar ====== */
    ::-webkit-scrollbar {
        width: 6px;
    }
    ::-webkit-scrollbar-track {
        background: transparent;
    }
    ::-webkit-scrollbar-thumb {
        background: rgba(148, 163, 184, 0.3);
        border-radius: 3px;
        backdrop-filter: blur(4px);
    }
    ::-webkit-scrollbar-thumb:hover {
        background: rgba(148, 163, 184, 0.5);
    }

    /* ====== Progress bar override ====== */
    .stProgress > div > div > div > div {
        background: linear-gradient(90deg, var(--primary-500), var(--accent-500)) !important;
        border-radius: 999px !important;
        box-shadow: 0 0 12px rgba(99, 102, 241, 0.3) !important;
    }

    /* ====== Download button - Glass Green ====== */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #059669 0%, #10B981 100%) !important;
        border: none !important;
        border-radius: var(--radius-md) !important;
        font-weight: 600 !important;
        color: #FFFFFF !important;
        box-shadow:
            0 0 15px rgba(16, 185, 129, 0.3),
            0 0 30px rgba(16, 185, 129, 0.15),
            0 8px 25px rgba(0, 0, 0, 0.12) !important;
        transition: all var(--transition-normal) !important;
        position: relative;
        overflow: hidden;
        text-shadow: 0 1px 2px rgba(0, 0, 0, 0.15);
    }
    .stDownloadButton > button::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 50%;
        background: linear-gradient(180deg, rgba(255, 255, 255, 0.15), transparent);
        pointer-events: none;
    }
    .stDownloadButton > button:hover {
        transform: translateY(-3px) !important;
        box-shadow:
            0 0 20px rgba(16, 185, 129, 0.45),
            0 0 40px rgba(16, 185, 129, 0.2),
            0 12px 35px rgba(0, 0, 0, 0.15) !important;
    }

    /* ====== Settings page section headings ====== */
    .settings-section-title {
        font-size: 1rem;
        font-weight: 600;
        color: var(--neutral-800);
        margin-bottom: 1rem;
        display: flex;
        align-items: center;
        gap: 0.5rem;
    }

    /* ====== Toggle override ====== */
    .stToggle > label {
        color: var(--neutral-700) !important;
    }

    /* ====== Global Glass Animations ====== */
    @keyframes glassShimmer {
        0% { background-position: -200% 0; }
        100% { background-position: 200% 0; }
    }

    /* ====== Selection highlight ====== */
    ::selection {
        background: rgba(99, 102, 241, 0.2);
        color: var(--neutral-900);
    }

    /* ====== SVG Icon System ====== */
    .svg-icon {
        display: inline-flex;
        align-items: center;
        justify-content: center;
        vertical-align: middle;
        line-height: 0;
    }
    .svg-icon svg {
        display: block;
    }

    /* Icon size variants */
    .svg-icon--xs svg { width: 14px; height: 14px; }
    .svg-icon--sm svg { width: 18px; height: 18px; }
    .svg-icon--md svg { width: 22px; height: 22px; }
    .svg-icon--lg svg { width: 28px; height: 28px; }
    .svg-icon--xl svg { width: 36px; height: 36px; }
    .svg-icon--2xl svg { width: 48px; height: 48px; }
    .svg-icon--3xl svg { width: 64px; height: 64px; }

    /* Icon glow effects */
    .icon-glow-primary {
        filter: drop-shadow(0 0 6px var(--glow-primary));
    }
    .icon-glow-accent {
        filter: drop-shadow(0 0 6px var(--glow-accent));
    }
    .icon-glow-success {
        filter: drop-shadow(0 0 6px var(--glow-success));
    }
    .icon-glow-error {
        filter: drop-shadow(0 0 6px rgba(239, 68, 68, 0.45));
    }
    .icon-glow-warning {
        filter: drop-shadow(0 0 6px rgba(245, 158, 11, 0.45));
    }
    .icon-glow-white {
        filter: drop-shadow(0 0 8px rgba(255, 255, 255, 0.5));
    }
    .icon-glow-soft {
        filter: drop-shadow(0 0 4px var(--glow-primary-soft));
    }

    /* Icon animations */
    @keyframes iconPulse {
        0%, 100% { opacity: 1; transform: scale(1); }
        50% { opacity: 0.8; transform: scale(0.95); }
    }
    @keyframes iconRotate {
        from { transform: rotate(0deg); }
        to { transform: rotate(360deg); }
    }
    @keyframes iconBounce {
        0%, 100% { transform: translateY(0); }
        50% { transform: translateY(-3px); }
    }
    @keyframes iconShimmer {
        0% { filter: drop-shadow(0 0 4px var(--glow-primary-soft)); }
        50% { filter: drop-shadow(0 0 10px var(--glow-primary)); }
        100% { filter: drop-shadow(0 0 4px var(--glow-primary-soft)); }
    }
    @keyframes rocketFlame {
        0%, 100% { opacity: 0.7; transform: scaleY(1); }
        50% { opacity: 1; transform: scaleY(1.15); }
    }

    .icon-pulse { animation: iconPulse 2s ease-in-out infinite; }
    .icon-rotate { animation: iconRotate 8s linear infinite; }
    .icon-bounce { animation: iconBounce 1.5s ease-in-out infinite; }
    .icon-shimmer { animation: iconShimmer 2.5s ease-in-out infinite; }

    /* Hero banner icon styling */
    .hero-icon {
        display: inline-flex;
        align-items: center;
        justify-content: center;
        margin-right: 0.35rem;
        vertical-align: middle;
        filter: drop-shadow(0 2px 8px rgba(255, 255, 255, 0.3));
    }
    .hero-icon svg {
        width: 36px;
        height: 36px;
    }

    /* Sidebar brand icon */
    .sidebar-brand-icon .svg-icon svg {
        width: 40px;
        height: 40px;
    }

    /* Icon in card header - inherits container styling */
    .card-header-icon .svg-icon svg {
        width: 20px;
        height: 20px;
    }

    /* Config item icon */
    .config-item-icon .svg-icon svg {
        width: 18px;
        height: 18px;
    }

    /* Stat icon */
    .stat-icon .svg-icon svg {
        width: 24px;
        height: 24px;
    }

    /* File item icon */
    .file-item-icon .svg-icon svg {
        width: 20px;
        height: 20px;
    }

    /* Button inline icon spacing */
    .btn-icon-left {
        display: inline-flex;
        align-items: center;
        gap: 0.4rem;
    }

    /* Empty state icon */
    .empty-state-icon .svg-icon svg {
        width: 56px;
        height: 56px;
    }

    /* Warning inline icon */
    .warning-icon-inline svg {
        width: 28px;
        height: 28px;
    }

    /* ====== Anti-flicker for sidebar ====== */
    [data-testid="stSidebar"] {
        will-change: transform;
        transform: translateZ(0);
    }
    [data-testid="stSidebar"] .sidebar-brand,
    [data-testid="stSidebar"] .tip-box,
    [data-testid="stSidebar"] .sidebar-section-label {
        will-change: opacity;
        opacity: 1;
        transition: opacity 0.1s ease-in-out;
    }
    [data-testid="stSidebar"] .stButton {
        will-change: transform;
        transform: translateZ(0);
    }
    [data-testid="stSidebar"] [data-testid="stVerticalBlock"] {
        will-change: contents;
    }

    /* ====== Page transition animation ====== */
    .stApp > header {
        display: none !important;
    }

    /* Main content area - fade in with upward slide */
    .block-container {
        animation: contentFadeIn 0.4s ease-out;
    }
    @keyframes contentFadeIn {
        from {
            opacity: 0;
            transform: translateY(10px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }

    /* Sidebar - slide in from left */
    [data-testid="stSidebar"] {
        animation: sidebarSlideIn 0.3s ease-out;
    }
    @keyframes sidebarSlideIn {
        from {
            transform: translateX(-20px);
            opacity: 0.8;
        }
        to {
            transform: translateX(0);
            opacity: 1;
        }
    }

    /* Card fade-in with upward slide */
    .card,
    .stat-card,
    .glass-card {
        animation: cardFadeIn 0.5s ease-out;
        animation-fill-mode: both;
    }
    @keyframes cardFadeIn {
        from {
            opacity: 0;
            transform: translateY(15px) scale(0.98);
        }
        to {
            opacity: 1;
            transform: translateY(0) scale(1);
        }
    }

    /* Staggered card animation delays */
    .card:nth-child(1) { animation-delay: 0.05s; }
    .card:nth-child(2) { animation-delay: 0.1s; }
    .card:nth-child(3) { animation-delay: 0.15s; }
    .card:nth-child(4) { animation-delay: 0.2s; }
    .card:nth-child(5) { animation-delay: 0.25s; }

    .stat-card:nth-child(1) { animation-delay: 0.1s; }
    .stat-card:nth-child(2) { animation-delay: 0.2s; }
    .stat-card:nth-child(3) { animation-delay: 0.3s; }

    /* Hero banner entrance */
    .hero-banner {
        animation: heroFadeIn 0.6s ease-out;
    }
    @keyframes heroFadeIn {
        from {
            opacity: 0;
            transform: translateY(8px) scale(0.99);
        }
        to {
            opacity: 1;
            transform: translateY(0) scale(1);
        }
    }

    /* Button press feedback */
    .stButton > button {
        transition: all 0.2s ease !important;
    }
    .stButton > button:active {
        transform: scale(0.97) !important;
    }

    /* Skeleton / shimmer loading effect */
    .skeleton {
        background: linear-gradient(90deg,
            rgba(148, 163, 184, 0.08) 25%,
            rgba(148, 163, 184, 0.15) 50%,
            rgba(148, 163, 184, 0.08) 75%
        );
        background-size: 200% 100%;
        animation: shimmer 1.5s infinite;
        border-radius: var(--radius-md);
    }
    [data-testid="stSidebar"] .skeleton {
        background: linear-gradient(90deg,
            rgba(255, 255, 255, 0.05) 25%,
            rgba(255, 255, 255, 0.1) 50%,
            rgba(255, 255, 255, 0.05) 75%
        );
        background-size: 200% 100%;
    }
    @keyframes shimmer {
        0% { background-position: 200% 0; }
        100% { background-position: -200% 0; }
    }

    /* Spinner overlay fade */
    .stSpinner {
        animation: spinnerFadeIn 0.2s ease-out;
    }
    @keyframes spinnerFadeIn {
        from { opacity: 0; }
        to { opacity: 1; }
    }

    /* Alert entrance */
    .stAlert {
        animation: alertSlideIn 0.3s ease-out;
    }
    @keyframes alertSlideIn {
        from {
            opacity: 0;
            transform: translateY(6px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
</style>
"""

st.markdown(DESIGN_SYSTEM_CSS, unsafe_allow_html=True)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _human_size(size_bytes: int) -> str:
    """Convert bytes to human-readable string."""
    if size_bytes < 1024:
        return f"{size_bytes} B"
    if size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    return f"{size_bytes / 1024 / 1024:.1f} MB"


def _file_icon(ext: str) -> str:
    """Return a styled SVG icon for a file extension."""
    icon_map = {
        ".py": "python",
        ".js": "javascript",
        ".ts": "typescript",
        ".tsx": "react",
        ".jsx": "react",
        ".html": "html",
        ".css": "css",
        ".json": "json",
        ".md": "markdown",
        ".txt": "text",
        ".docx": "word",
        ".pdf": "pdf",
        ".yaml": "config",
        ".yml": "config",
        ".sh": "shell",
        ".sql": "database",
        ".csv": "chart",
    }
    return _svg_icon(icon_map.get(ext, "file"), size="md")


def _svg_icon(name: str, size: str = "md", extra_class: str = "") -> str:
    """Return an inline SVG icon as an HTML string.

    Args:
        name: Icon name (e.g. 'doc', 'gear', 'rocket').
        size: One of 'xs', 'sm', 'md', 'lg', 'xl', '2xl', '3xl'.
        extra_class: Additional CSS classes to apply.

    Returns:
        HTML string with the SVG wrapped in a span.
    """
    cls = f"svg-icon svg-icon--{size}"
    if extra_class:
        cls += f" {extra_class}"

    icons = {
        # ── Document + pen (hero banner icon) ──
        "doc": f'''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="docG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#C7D2FE"/><stop offset="100%" stop-color="#FFFFFF"/>
            </linearGradient></defs>
            <path d="M14 2H6C4.9 2 4 2.9 4 4V20C4 21.1 4.9 22 6 22H18C19.1 22 20 21.1 20 20V8L14 2Z" stroke="url(#docG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M14 2V8H20" stroke="url(#docG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M12 18V12" stroke="url(#docG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M9 15L12 12L15 15" stroke="url(#docG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        # ── Document (for card headers, generic) ──
        "docDark": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="docDG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#8B5CF6"/>
            </linearGradient></defs>
            <path d="M14 2H6C4.9 2 4 2.9 4 4V20C4 21.1 4.9 22 6 22H18C19.1 22 20 21.1 20 20V8L14 2Z" stroke="url(#docDG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M14 2V8H20" stroke="url(#docDG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M8 13H16" stroke="url(#docDG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M8 17H13" stroke="url(#docDG)" stroke-width="1.5" stroke-linecap="round"/>
        </svg>''',

        # ── Folder (working directory) ──
        "folder": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="folG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#A78BFA"/>
            </linearGradient></defs>
            <path d="M22 19C22 20.1 21.1 21 20 21H4C2.9 21 2 20.1 2 19V5C2 3.9 2.9 3 4 3H9L11 6H20C21.1 6 22 6.9 22 8V19Z" stroke="url(#folG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M2 10H22" stroke="url(#folG)" stroke-width="1.2" stroke-linecap="round" opacity="0.5"/>
        </svg>''',

        # ── Calendar (scan days) ──
        "calendar": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="calG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <rect x="3" y="4" width="18" height="18" rx="3" stroke="url(#calG)" stroke-width="1.5"/>
            <path d="M16 2V6" stroke="url(#calG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M8 2V6" stroke="url(#calG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M3 10H21" stroke="url(#calG)" stroke-width="1.5" stroke-linecap="round"/>
            <circle cx="8" cy="15" r="1" fill="url(#calG)" opacity="0.7"/>
            <circle cx="12" cy="15" r="1" fill="url(#calG)" opacity="0.7"/>
            <circle cx="16" cy="15" r="1" fill="url(#calG)" opacity="0.7"/>
        </svg>''',

        # ── Search (scan button) ──
        "search": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="schG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#FFFFFF"/><stop offset="100%" stop-color="#C7D2FE"/>
            </linearGradient></defs>
            <circle cx="10.5" cy="10.5" r="7" stroke="url(#schG)" stroke-width="1.8"/>
            <path d="M15.5 15.5L21 21" stroke="url(#schG)" stroke-width="1.8" stroke-linecap="round"/>
            <circle cx="10.5" cy="10.5" r="3" stroke="url(#schG)" stroke-width="0.8" opacity="0.4"/>
        </svg>''',

        # ── Chart (scan stats header) ──
        "chart": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="chG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#8B5CF6"/>
            </linearGradient></defs>
            <rect x="3" y="12" width="4" height="9" rx="1" stroke="url(#chG)" stroke-width="1.5"/>
            <rect x="10" y="6" width="4" height="15" rx="1" stroke="url(#chG)" stroke-width="1.5"/>
            <rect x="17" y="3" width="4" height="18" rx="1" stroke="url(#chG)" stroke-width="1.5"/>
        </svg>''',

        # ── File (file count stat) ──
        "file": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="fiG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <path d="M14 2H6C4.9 2 4 2.9 4 4V20C4 21.1 4.9 22 6 22H18C19.1 22 20 21.1 20 20V8L14 2Z" stroke="url(#fiG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M14 2V8H20" stroke="url(#fiG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        # ── Clock (scan days stat) ──
        "clock": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="clG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <circle cx="12" cy="12" r="9" stroke="url(#clG)" stroke-width="1.5"/>
            <path d="M12 7V12L15 15" stroke="url(#clG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        # ── Storage (file size stat) ──
        "storage": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="stG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <rect x="3" y="3" width="18" height="6" rx="2" stroke="url(#stG)" stroke-width="1.5"/>
            <rect x="3" y="15" width="18" height="6" rx="2" stroke="url(#stG)" stroke-width="1.5"/>
            <circle cx="7" cy="6" r="1" fill="url(#stG)"/>
            <circle cx="7" cy="18" r="1" fill="url(#stG)"/>
            <path d="M3 9V15" stroke="url(#stG)" stroke-width="1.5"/>
            <path d="M21 9V15" stroke="url(#stG)" stroke-width="1.5"/>
        </svg>''',

        # ── Gear (settings) ──
        "gear": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="geG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#8B5CF6"/>
            </linearGradient></defs>
            <circle cx="12" cy="12" r="3" stroke="url(#geG)" stroke-width="1.5"/>
            <path d="M12 1v2M12 21v2M4.22 4.22l1.42 1.42M18.36 18.36l1.42 1.42M1 12h2M21 12h2M4.22 19.78l1.42-1.42M18.36 5.64l1.42-1.42" stroke="url(#geG)" stroke-width="1.5" stroke-linecap="round"/>
            <circle cx="12" cy="12" r="7" stroke="url(#geG)" stroke-width="1" opacity="0.3"/>
        </svg>''',

        # ── Chip / AI (model) ──
        "chip": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="cpG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#8B5CF6"/>
            </linearGradient></defs>
            <rect x="6" y="6" width="12" height="12" rx="2" stroke="url(#cpG)" stroke-width="1.5"/>
            <path d="M9 2V6M15 2V6M9 18V22M15 18V22M2 9H6M18 9H24M2 15H6M18 15H24" stroke="url(#cpG)" stroke-width="1.5" stroke-linecap="round"/>
            <circle cx="12" cy="12" r="2" stroke="url(#cpG)" stroke-width="1"/>
        </svg>''',

        # ── Doc stack (template) ──
        "docStack": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="dsG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#8B5CF6"/>
            </linearGradient></defs>
            <path d="M8 2H4C2.9 2 2 2.9 2 4V20C2 21.1 2.9 22 4 22H16C17.1 22 18 21.1 18 20V8L12 2" stroke="url(#dsG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M12 2V8H18" stroke="url(#dsG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M20 14V20C20 21.1 20.9 22 22 22" stroke="url(#dsG)" stroke-width="1.5" stroke-linecap="round" opacity="0.5"/>
            <path d="M20 14L22 12" stroke="url(#dsG)" stroke-width="1.5" stroke-linecap="round" opacity="0.5"/>
        </svg>''',

        # ── Rocket (generate) ──
        "rocket": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="rkG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#FFFFFF"/><stop offset="100%" stop-color="#C7D2FE"/>
            </linearGradient></defs>
            <path d="M4.5 16.5C3 18 2.5 21.5 2.5 21.5C2.5 21.5 6 21 7.5 19.5C8.32 18.68 8.32 17.32 7.5 16.5C6.68 15.68 5.32 15.68 4.5 16.5Z" stroke="url(#rkG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M14.5 4.5L12 2C12 2 7.5 6 7.5 12C7.5 14.5 8.5 16.5 10 18L14.5 15.5" stroke="url(#rkG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M14.5 4.5C17 4.5 20 5.5 21 7.5C22 9.5 22 14.5 19.5 17L14.5 15.5" stroke="url(#rkG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <circle cx="10.5" cy="10.5" r="1.5" stroke="url(#rkG)" stroke-width="1.2"/>
            <path d="M7 19L4.5 21.5" stroke="url(#rkG)" stroke-width="1" stroke-linecap="round" opacity="0.4"/>
        </svg>''',

        # ── Download (export) ──
        "download": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="dlG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#FFFFFF"/><stop offset="100%" stop-color="#A7F3D0"/>
            </linearGradient></defs>
            <path d="M21 15V19C21 20.1 20.1 21 19 21H5C3.9 21 3 20.1 3 19V15" stroke="url(#dlG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M12 3V15" stroke="url(#dlG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M7 10L12 15L17 10" stroke="url(#dlG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        # ── Check circle (success) ──
        "checkCircle": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="ccG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#10B981"/><stop offset="100%" stop-color="#34D399"/>
            </linearGradient></defs>
            <circle cx="12" cy="12" r="9" stroke="url(#ccG)" stroke-width="1.5"/>
            <path d="M8 12.5L10.5 15L16 9" stroke="url(#ccG)" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        # ── Warning triangle ──
        "warning": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="wnG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#F59E0B"/><stop offset="100%" stop-color="#FBBF24"/>
            </linearGradient></defs>
            <path d="M10.29 3.86L1.82 18C1.64 18.31 1.55 18.67 1.55 19.03C1.56 19.74 2.12 20.31 2.83 20.33H21.17C21.88 20.31 22.44 19.74 22.45 19.03C22.45 18.67 22.36 18.31 22.18 18L13.71 3.86C13.38 3.29 12.71 2.95 12 2.95C11.29 2.95 10.62 3.29 10.29 3.86Z" stroke="url(#wnG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M12 9V13" stroke="url(#wnG)" stroke-width="1.8" stroke-linecap="round"/>
            <circle cx="12" cy="17" r="0.8" fill="url(#wnG)"/>
        </svg>''',

        # ── X circle (error) ──
        "xCircle": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="xcG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#EF4444"/><stop offset="100%" stop-color="#F87171"/>
            </linearGradient></defs>
            <circle cx="12" cy="12" r="9" stroke="url(#xcG)" stroke-width="1.5"/>
            <path d="M9 9L15 15M15 9L9 15" stroke="url(#xcG)" stroke-width="1.8" stroke-linecap="round"/>
        </svg>''',

        # ── Info circle ──
        "info": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="inG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#3B82F6"/><stop offset="100%" stop-color="#60A5FA"/>
            </linearGradient></defs>
            <circle cx="12" cy="12" r="9" stroke="url(#inG)" stroke-width="1.5"/>
            <path d="M12 8V12" stroke="url(#inG)" stroke-width="1.8" stroke-linecap="round"/>
            <circle cx="12" cy="16" r="0.8" fill="url(#inG)"/>
        </svg>''',

        # ── Home (sidebar nav) ──
        "home": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="hmG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#C7D2FE"/><stop offset="100%" stop-color="#E0E7FF"/>
            </linearGradient></defs>
            <path d="M3 12L12 3L21 12" stroke="url(#hmG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M5 10V20C5 20.55 5.45 21 6 21H10V15H14V21H18C18.55 21 19 20.55 19 20V10" stroke="url(#hmG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        # ── Compass (sidebar brand) ──
        "compass": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="coG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#C7D2FE"/><stop offset="100%" stop-color="#FFFFFF"/>
            </linearGradient></defs>
            <circle cx="12" cy="12" r="9" stroke="url(#coG)" stroke-width="1.5"/>
            <path d="M16.24 7.76L14.12 14.12L7.76 16.24L9.88 9.88L16.24 7.76Z" stroke="url(#coG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <circle cx="12" cy="12" r="1" fill="url(#coG)" opacity="0.6"/>
        </svg>''',

        # ── Bolt / Lightning (concurrency) ──
        "bolt": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="btG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <path d="M13 2L3 14H12L11 22L21 10H12L13 2Z" stroke="url(#btG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        # ── Key (API key) ──
        "key": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="kyG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#8B5CF6"/>
            </linearGradient></defs>
            <circle cx="8" cy="15" r="5" stroke="url(#kyG)" stroke-width="1.5"/>
            <path d="M11.5 11.5L21 2" stroke="url(#kyG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M17 6L21 2" stroke="url(#kyG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M15 8L18 5" stroke="url(#kyG)" stroke-width="1.5" stroke-linecap="round"/>
        </svg>''',

        # ── Globe (URL) ──
        "globe": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="glG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <circle cx="12" cy="12" r="9" stroke="url(#glG)" stroke-width="1.5"/>
            <ellipse cx="12" cy="12" rx="4" ry="9" stroke="url(#glG)" stroke-width="1.2"/>
            <path d="M3 12H21" stroke="url(#glG)" stroke-width="1.2"/>
        </svg>''',

        # ── Package (model name) ──
        "package": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="pkG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#8B5CF6"/>
            </linearGradient></defs>
            <path d="M12 2L2 7V17L12 22L22 17V7L12 2Z" stroke="url(#pkG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M12 22V12" stroke="url(#pkG)" stroke-width="1.5"/>
            <path d="M22 7L12 12L2 7" stroke="url(#pkG)" stroke-width="1.5"/>
            <path d="M7 4.5L17 9.5" stroke="url(#pkG)" stroke-width="1" opacity="0.4"/>
        </svg>''',

        # ── Link (test connection) ──
        "link": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="lnG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <path d="M10 13C10.4 13.8 10.97 14.5 11.67 15.05C12.37 15.6 13.18 15.98 14.05 16.16C14.92 16.34 15.82 16.32 16.68 16.09C17.54 15.86 18.34 15.44 19 14.86L21 12.86" stroke="url(#lnG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M14 11C13.6 10.2 13.03 9.5 12.33 8.95C11.63 8.4 10.82 8.02 9.95 7.84C9.08 7.66 8.18 7.68 7.32 7.91C6.46 8.14 5.66 8.56 5 9.14L3 11.14" stroke="url(#lnG)" stroke-width="1.5" stroke-linecap="round"/>
        </svg>''',

        # ── Wrench (advanced settings) ──
        "wrench": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="wrG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <path d="M14.7 6.3C14.4 5.9 14 5.5 13.5 5.2C12.3 4.3 10.7 4.1 9.3 4.6C7.9 5.1 6.9 6.3 6.6 7.7C6.3 9.1 6.8 10.6 7.9 11.5L12.5 16.1C13.4 17.2 14.9 17.7 16.3 17.4C17.7 17.1 18.9 16.1 19.4 14.7C19.9 13.3 19.7 11.7 18.8 10.5C18.5 10 18.1 9.6 17.7 9.3L14.7 6.3Z" stroke="url(#wrG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M4 20L9 15" stroke="url(#wrG)" stroke-width="1.5" stroke-linecap="round"/>
        </svg>''',

        # ── Inbox (empty state) ──
        "inbox": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="ibG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#94A3B8"/><stop offset="100%" stop-color="#CBD5E1"/>
            </linearGradient></defs>
            <path d="M22 12H16L14 15H10L8 12H2" stroke="url(#ibG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M5.45 5.11L2 12V18C2 19.1 2.9 20 4 20H20C21.1 20 22 19.1 22 18V12L18.55 5.11C18.21 4.58 17.63 4.26 17.01 4.27H6.99C6.37 4.26 5.79 4.58 5.45 5.11Z" stroke="url(#ibG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        # ── Sparkle (AI badge) ──
        "sparkle": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="spG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#FBBF24"/><stop offset="100%" stop-color="#F59E0B"/>
            </linearGradient></defs>
            <path d="M12 2L13.09 8.26L18 6L14.74 10.91L21 12L14.74 13.09L18 18L13.09 15.74L12 22L10.91 15.74L6 18L9.26 13.09L3 12L9.26 10.91L6 6L10.91 8.26L12 2Z" stroke="url(#spG)" stroke-width="1.2" stroke-linejoin="round"/>
        </svg>''',

        # ── Save (save config) ──
        "save": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="svG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#FFFFFF"/><stop offset="100%" stop-color="#C7D2FE"/>
            </linearGradient></defs>
            <path d="M19 21H5C3.9 21 3 20.1 3 19V5C3 3.9 3.9 3 5 3H16L21 8V19C21 20.1 20.1 21 19 21Z" stroke="url(#svG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M17 21V13H7V21" stroke="url(#svG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M7 3V8H15" stroke="url(#svG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        # ── Log / List (log level) ──
        "log": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="loG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <path d="M8 6H21" stroke="url(#loG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M8 12H21" stroke="url(#loG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M8 18H21" stroke="url(#loG)" stroke-width="1.5" stroke-linecap="round"/>
            <circle cx="4" cy="6" r="1" fill="url(#loG)"/>
            <circle cx="4" cy="12" r="1" fill="url(#loG)"/>
            <circle cx="4" cy="18" r="1" fill="url(#loG)"/>
        </svg>''',

        # ── Cache / Database (cache toggle) ──
        "cache": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="caG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <ellipse cx="12" cy="5" rx="9" ry="3" stroke="url(#caG)" stroke-width="1.5"/>
            <path d="M3 5V19C3 20.66 7.03 22 12 22S21 20.66 21 19V5" stroke="url(#caG)" stroke-width="1.5"/>
            <path d="M3 12C3 13.66 7.03 15 12 15S21 13.66 21 12" stroke="url(#caG)" stroke-width="1.5"/>
        </svg>''',

        # ── Upload (file upload) ──
        "upload": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="upG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <path d="M21 15V19C21 20.1 20.1 21 19 21H5C3.9 21 3 20.1 3 19V15" stroke="url(#upG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M12 3V15" stroke="url(#upG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M7 8L12 3L17 8" stroke="url(#upG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        # ── Trash (delete) ──
        "trash": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="trG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#EF4444"/><stop offset="100%" stop-color="#F87171"/>
            </linearGradient></defs>
            <path d="M3 6H21" stroke="url(#trG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M19 6V20C19 21.1 18.1 22 17 22H7C5.9 22 5 21.1 5 20V6" stroke="url(#trG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M8 6V4C8 2.9 8.9 2 10 2H14C15.1 2 16 2.9 16 4V6" stroke="url(#trG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M10 11V17" stroke="url(#trG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M14 11V17" stroke="url(#trG)" stroke-width="1.5" stroke-linecap="round"/>
        </svg>''',

        # ── Arrow left (back) ──
        "arrowLeft": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="alG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#FFFFFF"/><stop offset="100%" stop-color="#C7D2FE"/>
            </linearGradient></defs>
            <path d="M19 12H5" stroke="url(#alG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M12 19L5 12L12 5" stroke="url(#alG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        # ── File type icons for _file_icon() ──
        "python": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="pyG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#3B82F6"/><stop offset="100%" stop-color="#60A5FA"/>
            </linearGradient></defs>
            <path d="M12 2C9 2 8 3 8 4V8H14V9H6C4 9 2 10.5 2 13.5S4 18 6 18H8V15C8 13.5 9 12 11 12H15C16.1 12 17 11.1 17 10V4C17 3 16 2 14 2H12Z" stroke="url(#pyG)" stroke-width="1.2"/>
            <circle cx="10" cy="5.5" r="0.8" fill="url(#pyG)"/>
            <path d="M12 22C15 22 16 21 16 20V16H10V15C10 15 12 15 14 15C16 15 18 13.5 18 10.5S16 6 14 6H12V9C12 10.5 11 12 9 12H5C3.9 12 3 12.9 3 14V20C3 21.1 3.9 22 5 22H12Z" stroke="url(#pyG)" stroke-width="1.2"/>
            <circle cx="14" cy="18.5" r="0.8" fill="url(#pyG)"/>
        </svg>''',

        "javascript": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="jsG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#FBBF24"/><stop offset="100%" stop-color="#F59E0B"/>
            </linearGradient></defs>
            <rect x="2" y="2" width="20" height="20" rx="3" stroke="url(#jsG)" stroke-width="1.3"/>
            <path d="M8 18C8 18 7 18 7 16.5V11" stroke="url(#jsG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M13 18C13 18 16 18 16 15C16 13 14 12 13 12C12 12 11 12.5 11 13.5" stroke="url(#jsG)" stroke-width="1.5" stroke-linecap="round"/>
        </svg>''',

        "typescript": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="tsG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#3B82F6"/><stop offset="100%" stop-color="#60A5FA"/>
            </linearGradient></defs>
            <rect x="2" y="2" width="20" height="20" rx="3" stroke="url(#tsG)" stroke-width="1.3"/>
            <path d="M11 7H7V18" stroke="url(#tsG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M7 12H10" stroke="url(#tsG)" stroke-width="1.5" stroke-linecap="round"/>
            <path d="M15 12C17 12 19 13 19 15C19 17 17 18 15 18H14V12" stroke="url(#tsG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M14 15H18" stroke="url(#tsG)" stroke-width="1.3" stroke-linecap="round"/>
        </svg>''',

        "react": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="rcG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#06B6D4"/><stop offset="100%" stop-color="#22D3EE"/>
            </linearGradient></defs>
            <ellipse cx="12" cy="12" rx="9" ry="4" stroke="url(#rcG)" stroke-width="1.3"/>
            <ellipse cx="12" cy="12" rx="9" ry="4" stroke="url(#rcG)" stroke-width="1.3" transform="rotate(60 12 12)"/>
            <ellipse cx="12" cy="12" rx="9" ry="4" stroke="url(#rcG)" stroke-width="1.3" transform="rotate(120 12 12)"/>
            <circle cx="12" cy="12" r="1.5" fill="url(#rcG)"/>
        </svg>''',

        "html": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="htG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#F97316"/><stop offset="100%" stop-color="#FB923C"/>
            </linearGradient></defs>
            <path d="M4 3L5.5 18L12 21L18.5 18L20 3H4Z" stroke="url(#htG)" stroke-width="1.3" stroke-linejoin="round"/>
            <path d="M8 7H16L15.5 13H9L8.5 17L12 18L15.5 17" stroke="url(#htG)" stroke-width="1.2" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        "css": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="csG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#3B82F6"/><stop offset="100%" stop-color="#60A5FA"/>
            </linearGradient></defs>
            <path d="M4 3L5.5 18L12 21L18.5 18L20 3H4Z" stroke="url(#csG)" stroke-width="1.3" stroke-linejoin="round"/>
            <path d="M8 7H16L15.5 13H9" stroke="url(#csG)" stroke-width="1.2" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M10 15L8.5 17L12 18L15.5 17" stroke="url(#csG)" stroke-width="1.2" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        "json": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="jG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#FBBF24"/><stop offset="100%" stop-color="#F59E0B"/>
            </linearGradient></defs>
            <path d="M8 3C5 3 4 4.5 4 6V9C4 10.1 3.1 11 2 11C3.1 11 4 11.9 4 13V18C4 19.5 5 21 8 21" stroke="url(#jG)" stroke-width="1.3" stroke-linecap="round"/>
            <path d="M16 3C19 3 20 4.5 20 6V9C20 10.1 20.9 11 22 11C20.9 11 20 11.9 20 13V18C20 19.5 19 21 16 21" stroke="url(#jG)" stroke-width="1.3" stroke-linecap="round"/>
        </svg>''',

        "markdown": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="mdG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <rect x="2" y="4" width="20" height="16" rx="2" stroke="url(#mdG)" stroke-width="1.3"/>
            <path d="M6 15V9L9 12L12 9V15" stroke="url(#mdG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M18 15V9L15 12" stroke="url(#mdG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        "text": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="txG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#64748B"/><stop offset="100%" stop-color="#94A3B8"/>
            </linearGradient></defs>
            <path d="M14 2H6C4.9 2 4 2.9 4 4V20C4 21.1 4.9 22 6 22H18C19.1 22 20 21.1 20 20V8L14 2Z" stroke="url(#txG)" stroke-width="1.3" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M14 2V8H20" stroke="url(#txG)" stroke-width="1.3" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M8 13H16" stroke="url(#txG)" stroke-width="1.2" stroke-linecap="round"/>
            <path d="M8 17H14" stroke="url(#txG)" stroke-width="1.2" stroke-linecap="round"/>
        </svg>''',

        "word": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="woG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#3B82F6"/><stop offset="100%" stop-color="#60A5FA"/>
            </linearGradient></defs>
            <path d="M14 2H6C4.9 2 4 2.9 4 4V20C4 21.1 4.9 22 6 22H18C19.1 22 20 21.1 20 20V8L14 2Z" stroke="url(#woG)" stroke-width="1.3" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M14 2V8H20" stroke="url(#woG)" stroke-width="1.3" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M8 13L10 18L12 13L14 18L16 13" stroke="url(#woG)" stroke-width="1.3" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        "pdf": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="pfG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#EF4444"/><stop offset="100%" stop-color="#F87171"/>
            </linearGradient></defs>
            <path d="M14 2H6C4.9 2 4 2.9 4 4V20C4 21.1 4.9 22 6 22H18C19.1 22 20 21.1 20 20V8L14 2Z" stroke="url(#pfG)" stroke-width="1.3" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M14 2V8H20" stroke="url(#pfG)" stroke-width="1.3" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M8 14H10C11.1 14 12 13.1 12 12V10.5C12 9.4 11.1 8.5 10 8.5H8V14Z" stroke="url(#pfG)" stroke-width="1.2" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M14 8.5V14" stroke="url(#pfG)" stroke-width="1.3" stroke-linecap="round"/>
            <path d="M16 8.5V14" stroke="url(#pfG)" stroke-width="1.3" stroke-linecap="round"/>
        </svg>''',

        "config": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="cfG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <path d="M14 2H6C4.9 2 4 2.9 4 4V20C4 21.1 4.9 22 6 22H18C19.1 22 20 21.1 20 20V8L14 2Z" stroke="url(#cfG)" stroke-width="1.3" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M14 2V8H20" stroke="url(#cfG)" stroke-width="1.3" stroke-linecap="round" stroke-linejoin="round"/>
            <circle cx="12" cy="14" r="2.5" stroke="url(#cfG)" stroke-width="1.2"/>
            <path d="M12 10V11.5" stroke="url(#cfG)" stroke-width="1.2" stroke-linecap="round"/>
            <path d="M12 16.5V18" stroke="url(#cfG)" stroke-width="1.2" stroke-linecap="round"/>
            <path d="M9.5 14H8" stroke="url(#cfG)" stroke-width="1.2" stroke-linecap="round"/>
            <path d="M16 14H14.5" stroke="url(#cfG)" stroke-width="1.2" stroke-linecap="round"/>
        </svg>''',

        "shell": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="shG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#10B981"/><stop offset="100%" stop-color="#34D399"/>
            </linearGradient></defs>
            <rect x="2" y="3" width="20" height="18" rx="3" stroke="url(#shG)" stroke-width="1.3"/>
            <path d="M6 10L9.5 13L6 16" stroke="url(#shG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M12 16H18" stroke="url(#shG)" stroke-width="1.5" stroke-linecap="round"/>
        </svg>''',

        "database": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="dbG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <ellipse cx="12" cy="5" rx="8" ry="3" stroke="url(#dbG)" stroke-width="1.3"/>
            <path d="M4 5V19C4 20.66 7.58 22 12 22S20 20.66 20 19V5" stroke="url(#dbG)" stroke-width="1.3"/>
            <path d="M4 12C4 13.66 7.58 15 12 15S20 13.66 20 12" stroke="url(#dbG)" stroke-width="1.3"/>
        </svg>''',

        # ── Shield / OAuth lock ──
        "shield": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="shdG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <path d="M12 2L3 7V12C3 17.55 7.84 22.74 12 24C16.16 22.74 21 17.55 21 12V7L12 2Z" stroke="url(#shdG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M8 12L10.5 14.5L16 9" stroke="url(#shdG)" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>''',

        # ── External link ──
        "externalLink": '''<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <defs><linearGradient id="elG" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#6366F1"/><stop offset="100%" stop-color="#818CF8"/>
            </linearGradient></defs>
            <path d="M18 13V19C18 20.1 17.1 21 16 21H5C3.9 21 3 20.1 3 19V8C3 6.9 3.9 6 5 6H11" stroke="url(#elG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M15 3H21V9" stroke="url(#elG)" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M10 14L21 3" stroke="url(#elG)" stroke-width="1.5" stroke-linecap="round"/>
        </svg>''',
    }

    svg = icons.get(name, icons["file"])
    return f'<span class="{cls}">{svg}</span>'


def _get_provider_choices() -> Tuple[List[str], Dict[str, str]]:
    """Return (display_names_list, display_to_key_mapping)."""
    from llm.factory import get_available_providers
    
    providers_info = get_available_providers()
    if not providers_info:
        return [], {}

    names = []
    mapping = {}
    for key, info in providers_info.items():
        if info['configured'] == 'yes':
            # 直接显示模型名称
            model_name = info.get('model', '')
            if model_name:
                display = model_name
            else:
                display = info['display_name']
            names.append(display)
            mapping[display] = key
    return names, mapping


def _md_to_txt(md_text: str) -> str:
    """Convert Markdown to plain text by stripping formatting."""
    text = re.sub(r"^#{1,6}\s+", "", md_text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"~~(.+?)~~", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"^[-*+]\s+", "• ", text, flags=re.MULTILINE)
    text = re.sub(r"^\d+\.\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^---+$", "─" * 40, text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _md_to_docx(md_text: str) -> bytes:
    """Convert Markdown text to a Word document (.docx)."""
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    lines = md_text.split("\n")
    in_code_block = False
    code_lines: List[str] = []
    table_rows: List[List[str]] = []
    in_table = False

    for line in lines:
        if line.strip().startswith("```"):
            if in_code_block:
                if code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run("\n".join(code_lines))
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                code_lines = []
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        stripped = line.strip()

        if not stripped:
            # If we were in a table, flush it
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            continue

        # Handle table rows
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            # Check if it's a separator row
            if set(stripped.replace("|", "").replace("-", "").replace(":", "").strip()) == set():
                continue  # Skip separator rows
            # Parse cells
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            table_rows.append(cells)
            continue
        
        # If we were in a table and hit non-table line, flush table
        if in_table and table_rows:
            _add_table_to_doc(doc, table_rows)
            table_rows = []
            in_table = False

        if stripped.startswith("# "):
            h = doc.add_heading(stripped[2:].strip(), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
        elif stripped.startswith("---"):
            p = doc.add_paragraph()
            p.add_run("─" * 50)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped[2:])
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            # 处理来源链接
            text = _process_source_links_for_docx(text)
            doc.add_paragraph(text, style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            text = re.sub(r"\d+\.\s", "", text, count=1)
            # 处理来源链接
            text = _process_source_links_for_docx(text)
            doc.add_paragraph(text, style="List Number")
        else:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            # 处理来源链接
            text = _process_source_links_for_docx(text)
            doc.add_paragraph(text)

    # Flush any remaining table
    if in_table and table_rows:
        _add_table_to_doc(doc, table_rows)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _process_source_links_for_docx(text: str) -> str:
    """处理 Markdown 文本中的来源链接，使其在 Word 文档中正确显示。

    将 [path](file:///path) 格式的链接转换为纯文本显示。
    """
    import re
    # 将 [text](file:///path) 转换为 "text"
    text = re.sub(r'\[([^\]]+)\]\(file:///[^\)]+\)', r'\1', text)
    return text


def _add_table_to_doc(doc, rows: List[List[str]]) -> None:
    """Add a markdown table to the Word document."""
    from docx.shared import Pt, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    if not rows:
        return
    
    # Determine number of columns
    num_cols = max(len(row) for row in rows)
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Fill table
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                # Make header row bold
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    doc.add_paragraph()  # Add space after table


def _md_to_pdf(md_text: str) -> bytes:
    """Convert Markdown text to PDF using fpdf2 with Chinese font support."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)
    pdf.add_page()

    # Try to load a Chinese font
    font_loaded = False
    font_paths = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\msyhbd.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\simfang.ttf",
    ]
    for font_path in font_paths:
        if Path(font_path).exists():
            try:
                pdf.add_font("Chinese", "", font_path, uni=True)
                pdf.set_font("Chinese", size=10)
                font_loaded = True
                break
            except Exception:
                continue

    if not font_loaded:
        pdf.set_font("Helvetica", size=10)

    def write_text(text: str, size: int = 10, bold: bool = False):
        """Helper to write text with proper font settings."""
        if font_loaded:
            pdf.set_font("Chinese", size=size)
        else:
            style = "B" if bold else ""
            pdf.set_font("Helvetica", style, size=size)
        # Ensure we're at left margin
        pdf.set_x(15)
        pdf.multi_cell(0, size * 0.5, text)

    lines = md_text.split("\n")
    in_code_block = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            if font_loaded:
                pdf.set_font("Chinese", size=8)
            else:
                pdf.set_font("Courier", size=8)
            pdf.set_x(15)
            pdf.multi_cell(0, 4, stripped)
            if font_loaded:
                pdf.set_font("Chinese", size=10)
            else:
                pdf.set_font("Helvetica", size=10)
            continue

        if not stripped:
            pdf.ln(2)
            continue

        # Remove markdown formatting for PDF
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        
        # 处理来源链接 - 将 [text](file:///path) 转换为纯文本
        text = re.sub(r'\[([^\]]+)\]\(file:///[^\)]+\)', r'\1', text)

        if stripped.startswith("# "):
            write_text(text[2:].strip(), size=16, bold=True)
            pdf.ln(2)
        elif stripped.startswith("## "):
            write_text(text[3:].strip(), size=13, bold=True)
            pdf.ln(1)
        elif stripped.startswith("### "):
            write_text(text[4:].strip(), size=11, bold=True)
            pdf.ln(1)
        elif stripped.startswith("#### "):
            write_text(text[5:].strip(), size=10, bold=True)
        elif stripped.startswith("---"):
            pdf.line(15, pdf.get_y(), 195, pdf.get_y())
            pdf.ln(3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            content = text[2:].strip()
            write_text(f"  \u2022 {content}", size=10)
        elif re.match(r"^\d+\.\s", stripped):
            write_text(f"  {text}", size=10)
        else:
            write_text(text, size=10)

    return bytes(pdf.output())


def _get_export_data(report: str, fmt: str, days: int = 7) -> Tuple[bytes, str, str]:
    """Get export data for the selected format.

    Returns:
        (data_bytes, filename, mime_type)
    """
    week_range = get_week_range(days=days)
    base_name = f"周报_{week_range.replace(' ', '').replace('-', '_')}"

    if fmt == "Markdown (.md)":
        return report.encode("utf-8"), f"{base_name}.md", "text/markdown"
    elif fmt == "纯文本 (.txt)":
        return _md_to_txt(report).encode("utf-8"), f"{base_name}.txt", "text/plain"
    elif fmt == "Word (.docx)":
        try:
            return _md_to_docx(report), f"{base_name}.docx", (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Word导出失败: {e}")
            return report.encode("utf-8"), f"{base_name}.md", "text/markdown"
    elif fmt == "PDF (.pdf)":
        try:
            pdf_data = _md_to_pdf(report)
            return pdf_data, f"{base_name}.pdf", "application/pdf"
        except Exception as e:
            st.error(f"PDF导出失败: {e}")
            return report.encode("utf-8"), f"{base_name}.md", "text/markdown"
    else:
        return report.encode("utf-8"), f"{base_name}.md", "text/markdown"


def _get_model_display_name() -> str:
    """Get display name of current model config."""
    config_mode = st.session_state.get("config_mode", "使用 .env 配置")
    if config_mode == "手动配置":
        model = st.session_state.get("manual_model_name", "")
        return model if model else "未配置"
    else:
        provider_key = st.session_state.get("selected_provider_key")
        if provider_key:
            # 获取模型准确名称
            from llm.factory import get_available_providers
            providers_info = get_available_providers()
            if provider_key in providers_info:
                model_name = providers_info[provider_key].get('model', '')
                if model_name:
                    return model_name
            return PROVIDER_DISPLAY_NAMES.get(provider_key, provider_key)
        return "未检测到"


def _get_template_display_name() -> str:
    """Get display name of current template."""
    tpl_key = st.session_state.get("selected_template")
    if not tpl_key:
        return "标准模板"
    templates = list_templates()
    if tpl_key in templates:
        return templates[tpl_key]["name"]
    return tpl_key


# ---------------------------------------------------------------------------
# Data source helpers
# ---------------------------------------------------------------------------

# Canonical mapping: display name -> source key
DATASOURCE_OPTIONS = {
    "📁 本地文件": "file",
}


def _create_datasource(source_key: str):
    """Create a DataSource instance based on session_state config.

    Returns the DataSource instance, or ``None`` for the ``"file"`` source
    (which is handled separately via ``scan_folder``).
    """
    if source_key == "file":
        return None  # Handled by existing scan_folder flow

    return None


def _fetch_all_data(selected_sources: List[str], days: int) -> List[Dict]:
    """Fetch work items from all selected non-file data sources.

    The ``"file"`` source is handled separately via ``scan_folder`` in the
    main generation flow.  This function only fetches from API-based sources.

    Parameters
    ----------
    selected_sources : list[str]
        Display names of selected sources (e.g. ``["📁 本地文件", "🔀 Git提交记录"]``).
    days : int
        Number of days of history to retrieve.

    Returns
    -------
    list[dict]
        Combined list of standardised items from all non-file sources.
    """
    all_items: List[Dict] = []

    for source_name in selected_sources:
        source_key = DATASOURCE_OPTIONS.get(source_name)
        if source_key is None or source_key == "file":
            continue

        source = _create_datasource(source_key)
        if source is None:
            logger.warning("Skipping %s: configuration incomplete", source_name)
            continue

        if not source.is_available():
            logger.warning("Skipping %s: source not available", source_name)
            st.warning(f"⚠️ {source_name} 不可用，已跳过")
            continue

        try:
            items = source.fetch(days=days)
            logger.info("Fetched %d items from %s", len(items), source_name)
            all_items.extend(items)
        except Exception as exc:
            logger.error("Failed to fetch from %s: %s", source_name, exc)
            st.warning(f"⚠️ {source_name} 获取数据失败: {exc}")

    return all_items


# ---------------------------------------------------------------------------
# OAuth / Connection-test helpers
# ---------------------------------------------------------------------------

# Auth type classification
NO_AUTH_SOURCES = {"file"}

# Human-readable auth type labels
AUTH_TYPE_LABELS = {
    "file": "无需授权",
}


def _test_datasource_connection(source_key: str) -> Tuple[bool, str]:
    """Test connection for a given data source.

    Returns (success: bool, message: str).
    """
    try:
        source = _create_datasource(source_key)
        if source is None:
            return False, "配置不完整，请填写所有必填字段"
        if source.is_available():
            return True, "连接成功"
        return False, "数据源不可用，请检查配置"
    except ImportError as e:
        return False, f"缺少依赖: {e}"
    except Exception as e:
        return False, f"连接失败: {e}"


def _get_sensitivity_warnings(selected_sources: List[str]) -> List[str]:
    """Collect sensitivity warnings from all selected data sources.

    Returns
    -------
    list[str]
        Non-empty warning strings.
    """
    warnings: List[str] = []
    for source_name in selected_sources:
        source_key = DATASOURCE_OPTIONS.get(source_name)
        if source_key is None or source_key == "file":
            continue

        source = _create_datasource(source_key)
        if source is None:
            continue

        try:
            warning = source.get_sensitivity_warning()
            if warning:
                warnings.append(f"{source_name}: {warning}")
        except Exception:
            pass

    return warnings


def _convert_items_to_analyses(items: List[Dict]) -> List[Dict]:
    """Convert data-source items into the analysis format expected by the generator.

    Each data-source item has keys: source_type, title, content, metadata, timestamp.
    The generator expects dicts with: file_name, file_path, analysis, status.

    Parameters
    ----------
    items : list[dict]
        Standardised items from data sources.

    Returns
    -------
    list[dict]
        Analysis-format dicts compatible with ``generate_report``.
    """
    analyses: List[Dict] = []
    for item in items:
        source_type = item.get("source_type", "unknown")
        title = item.get("title", "(无标题)")
        content = item.get("content", "")
        timestamp = item.get("timestamp", "")

        # Build analysis text from the item's content
        analysis_parts = []
        if content:
            analysis_parts.append(content)
        if timestamp:
            analysis_parts.append(f"时间: {timestamp}")

        analysis_text = "\n".join(analysis_parts)

        analyses.append({
            "file_name": f"[{source_type}] {title}",
            "file_path": source_type,
            "analysis": analysis_text,
            "status": "success",
        })

    return analyses


# ---------------------------------------------------------------------------
# Session state defaults
# ---------------------------------------------------------------------------

if "report" not in st.session_state:
    st.session_state.report = None
if "files" not in st.session_state:
    st.session_state.files = None
if "scan_done" not in st.session_state:
    st.session_state.scan_done = False
if "config_mode" not in st.session_state:
    st.session_state.config_mode = "使用 .env 配置"
if "manual_base_url" not in st.session_state:
    st.session_state.manual_base_url = ""
if "manual_api_key" not in st.session_state:
    st.session_state.manual_api_key = ""
if "manual_model_name" not in st.session_state:
    st.session_state.manual_model_name = ""
if "selected_template" not in st.session_state:
    st.session_state.selected_template = None
if "selected_provider_key" not in st.session_state:
    st.session_state.selected_provider_key = None
if "scan_days" not in st.session_state:
    st.session_state.scan_days = 7
if "work_dir" not in st.session_state:
    default_dir = Config.SCAN_DIRS.split(",")[0].strip()
    if default_dir == ".":
        default_dir = str(Path(__file__).resolve().parent)
    st.session_state.work_dir = default_dir
if "concurrency" not in st.session_state:
    st.session_state.concurrency = 3
if "use_cache" not in st.session_state:
    st.session_state.use_cache = True
if "log_level" not in st.session_state:
    st.session_state.log_level = "INFO"
if "selected_sources" not in st.session_state:
    st.session_state.selected_sources = ["📁 本地文件"]

# OAuth / Connection status tracking
if "source_connection_status" not in st.session_state:
    st.session_state.source_connection_status = {}  # {source_key: {"ok": bool, "msg": str}}

# 时间线/历史记录相关
if "file_changes" not in st.session_state:
    st.session_state.file_changes = None  # 文件变更信息
if "show_timeline" not in st.session_state:
    st.session_state.show_timeline = False  # 是否显示时间线
if "save_snapshot_on_generate" not in st.session_state:
    st.session_state.save_snapshot_on_generate = True  # 生成报告时自动保存快照

# ---------------------------------------------------------------------------
# Auto-detect available AI models at startup
# ---------------------------------------------------------------------------
if st.session_state.get("selected_provider_key") is None:
    try:
        _providers, _mapping = _get_provider_choices()
        if _providers:
            st.session_state.selected_provider_key = _mapping[_providers[0]]
    except Exception:
        pass  # Silently skip if detection fails

# ---------------------------------------------------------------------------
# Auto-scan callback
# ---------------------------------------------------------------------------

def auto_scan_callback():
    """自动扫描回调函数 - 在后台线程中执行扫描、分析、生成报告并保存快照.
    从 .scheduler_config.json 读取配置，不依赖 st.session_state。"""
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Auto scan callback started")
    # 从配置文件读取扫描配置（后台线程无法访问 st.session_state）
    scheduler = AutoScheduler()
    scan_cfg = scheduler.load_scan_config()

    work_dir = scan_cfg.get("work_dir", "")
    scan_days = scan_cfg.get("scan_days", 7)
    selected_template = scan_cfg.get("selected_template")
    config_mode = scan_cfg.get("config_mode", "使用 .env 配置")
    selected_provider_key = scan_cfg.get("selected_provider_key")
    retention_days = scan_cfg.get("retention_days", 30)

    if not work_dir:
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Auto scan skipped: no work_dir configured")
        logger.warning("自动扫描跳过: 未设置工作目录")
        return

    try:
        # 1. 扫描文件
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Auto scan: scanning {work_dir} for last {scan_days} days")
        files = scan_folder(work_dir, days=scan_days)
        if not files:
            print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Auto scan completed: no files found")
            logger.info("自动扫描完成: 未发现文件")
            return

        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Auto scan: found {len(files)} files")
        logger.info("自动扫描: 发现 %d 个文件", len(files))

        # 2. 创建 LLM provider
        provider = None
        if config_mode == "手动配置":
            base_url = scan_cfg.get("manual_base_url", "")
            api_key = scan_cfg.get("manual_api_key", "")
            model_name = scan_cfg.get("manual_model_name", "")
            if base_url and model_name:
                provider = CustomProvider(
                    api_key=api_key,
                    model=model_name,
                    base_url=base_url,
                )
        else:
            if selected_provider_key:
                provider = create_provider(selected_provider_key)

        if provider is None:
            print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Auto scan skipped: no AI model configured")
            logger.warning("自动扫描跳过: 未配置 AI 模型")
            return

        # 3. 获取历史快照内容用于增量分析
        _snapshot = get_latest_snapshot()
        _old_file_contents = get_snapshot_file_contents(_snapshot)

        # 4. 逐文件分析
        analyses = []
        for file_info in files:
            file_path = file_info.get("path", "")
            old_content = _old_file_contents.get(file_path)
            new_content = read_file(file_path)
            if new_content.startswith("[Error]") or new_content.startswith("[Skipped]"):
                old_content = None
            result = analyze_changes(file_info, provider, old_content, new_content)
            analyses.append(result)

        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Auto scan: analysis completed for {len(analyses)} files")
        logger.info("自动扫描: 分析完成 %d 个文件", len(analyses))

        # 5. 生成报告
        week_range = get_week_range(days=scan_days)
        previous_report = None
        latest_snapshot = get_latest_snapshot()
        if latest_snapshot and latest_snapshot.get("report"):
            previous_report = latest_snapshot["report"]

        report = generate_report(
            analyses,
            provider,
            week_range=week_range,
            template_name=selected_template,
            previous_report=previous_report,
            file_changes=compare_with_latest(files, scan_days),
        )

        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Auto scan: report generated")
        logger.info("自动扫描: 报告生成完成")

        # 6. 保存快照
        snapshot_id = save_snapshot(
            files=files,
            analyses=analyses,
            report=report,
            metadata={
                "scan_days": scan_days,
                "work_dir": work_dir,
                "template": selected_template,
                "retention_days": retention_days,
                "auto_generated": True,
            }
        )
        if snapshot_id:
            print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Auto scan: snapshot saved {snapshot_id}")
            logger.info("自动扫描: 快照已保存 %s", snapshot_id)

        # 7. 归档周报到历史记录
        archive_id = archive_report(
            report_content=report,
            date_range=week_range,
            template_name=selected_template,
            source_files=files,
            work_dir=work_dir,
            scan_days=scan_days,
            metadata={"auto_generated": True},
        )
        if archive_id:
            print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Auto scan: report archived {archive_id}")
            logger.info("自动扫描: 周报已归档 %s", archive_id)

        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Auto scan all steps completed successfully")

    except Exception as e:
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Auto scan error: {e}")
        logger.error("自动扫描出错: %s", e, exc_info=True)
        import traceback
        traceback.print_exc()


# 初始化调度器
if "auto_scheduler" not in st.session_state:
    st.session_state.auto_scheduler = AutoScheduler()

# 从配置文件加载自动扫描设置到 session_state（确保刷新页面后设置不丢失）
_scheduler = st.session_state.auto_scheduler
if "auto_scan_enabled" not in st.session_state:
    st.session_state.auto_scan_enabled = _scheduler.enabled
    # 解析时间字符串
    try:
        _hour, _minute = _scheduler.time_str.split(":")
        st.session_state.auto_scan_hour = int(_hour)
        st.session_state.auto_scan_minute = int(_minute)
    except (ValueError, AttributeError):
        st.session_state.auto_scan_hour = 9
        st.session_state.auto_scan_minute = 0
    print(f"[Scheduler Init] Loaded from config: enabled={_scheduler.enabled}, time={_scheduler.time_str}")

# 应用启动时，如果调度器配置为启用则自动启动
if _scheduler.enabled and not _scheduler.running:
    _scheduler.set_callback(auto_scan_callback)
    _scheduler.start()
    print(f"[Scheduler Init] Scheduler auto-started, scheduled time: {_scheduler.time_str}")


# ---------------------------------------------------------------------------
# Page: Main (主页)
# ---------------------------------------------------------------------------

def main_page():
    """主页面 - 核心操作入口."""

    # -- Hero Banner --
    _hero_doc = _svg_icon("doc", size="lg", extra_class="hero-icon icon-glow-white")
    st.markdown(
        f"""
        <div class="hero-banner">
            <div class="hero-badge">{_svg_icon("sparkle", size="xs")} AI 驱动</div>
            <h1>{_hero_doc} 周报生成器</h1>
            <p>智能扫描工作文件，一键生成专业工作汇总</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # -- Working directory & Scan settings card --
    st.markdown(
        f"""
        <div class="card">
            <div class="card-header">
                <span class="card-header-icon">{_svg_icon("folder", size="md")}</span>
                <span class="card-header-title">工作目录</span>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    work_dir = st.text_input(
        "工作目录",
        value=st.session_state.work_dir,
        placeholder="输入要扫描的文件夹路径，例如 C:\\Users\\work\\projects",
        help="输入要扫描的文件夹绝对路径或相对路径",
        label_visibility="collapsed",
    )
    st.session_state.work_dir = work_dir

    col_slider, col_scan_btn = st.columns([3, 1])
    with col_slider:
        scan_days = st.number_input(
            "📅 扫描天数",
            min_value=1,
            max_value=30,
            value=st.session_state.scan_days,
            help="扫描最近N天内修改过的文件",
            key="scan_days_widget",
            on_change=lambda: setattr(st.session_state, 'scan_days', st.session_state.scan_days_widget),
        )
    with col_scan_btn:
        st.markdown("<div style='padding-top: 1.5rem'></div>", unsafe_allow_html=True)
        if st.button("🔍 扫描文件", use_container_width=True, type="primary"):
            if not work_dir or not work_dir.strip():
                st.error("请输入有效的工作目录路径")
            else:
                resolved = Path(work_dir).resolve()
                if not resolved.is_dir():
                    st.error(f"目录不存在: {resolved}")
                else:
                    with st.spinner("正在扫描文件..."):
                        # 获取所有文件（用于显示）
                        all_files = scan_folder_all(work_dir)
                        # 获取待分析文件（仅最近N天修改的）
                        files_to_analyze = scan_folder(work_dir, days=scan_days)
                    
                    st.session_state.files = all_files  # 显示所有文件
                    st.session_state.files_to_analyze = files_to_analyze  # 待分析文件
                    st.session_state.scan_done = True
                    st.session_state.report = None  # Reset report on new scan
                    
                    # 先进行变更检测（与历史快照比较，使用待分析文件）
                    changes = compare_with_latest(files_to_analyze, scan_days)
                    st.session_state.file_changes = changes
                    
                    # 再保存当前扫描快照（用于下次对比，保存待分析文件）
                    save_snapshot(
                        files=files_to_analyze,
                        analyses=[],
                        report=None,
                        metadata={
                            "scan_days": scan_days,
                            "work_dir": work_dir,
                            "scan_only": True,
                            "retention_days": st.session_state.get("retention_days", 30),
                        }
                    )
                    
                    # 变更信息在下方扫描统计区域统一显示，此处不重复显示

    # Close the card div (the card-header was opened above, content is outside)
    # We need to close the card properly - using container border instead

    # -- Stats & File list (after scan) --
    files = st.session_state.files
    files_to_analyze = st.session_state.get("files_to_analyze", [])

    if st.session_state.scan_done and files is not None:
        total_files = len(files)
        total_size = sum(f["size"] for f in files)
        analyze_count = len(files_to_analyze)

        # 获取变更信息
        changes = st.session_state.file_changes or {}
        added_count = len(changes.get("added", []))
        modified_count = len(changes.get("modified", []))
        removed_count = len(changes.get("removed", []))
        has_previous = changes.get("has_previous", False)
        found_match = changes.get("found_match", False)
        nearest_snapshot_time = changes.get("nearest_snapshot_time", "")

        st.markdown(
            f"""
            <div class="card">
                <div class="card-header">
                    <span class="card-header-icon">{_svg_icon("chart", size="md")}</span>
                    <span class="card-header-title">扫描统计</span>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        
        # 使用 Streamlit 原生组件显示统计信息
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("📄 总文件数", total_files)
        with col2:
            st.metric("📊 待分析", analyze_count)
        with col3:
            st.metric("📅 扫描天数", scan_days)
        with col4:
            st.metric("💾 总大小", _human_size(total_size))
        
        # 如果有变更，显示变更统计
        if has_previous and found_match:
            if added_count or modified_count:
                st.info(f"📊 变更检测（对比 {scan_days} 天前）：新增 {added_count} 个，修改 {modified_count} 个文件")
            if not added_count and not modified_count:
                st.info(f"📊 与 {scan_days} 天前相比无变化")
        elif has_previous and not found_match:
            st.warning(f"⚠️ 未找到 {scan_days} 天前的历史记录，所有文件将作为新增内容处理")
        elif not has_previous:
            st.info("📝 首次扫描，已保存为基准记录")

        # -- File list (expandable) --
        if total_files > 0:
            # 构建变更文件集合
            added_files = {f["path"] for f in changes.get("added", [])}
            modified_files = {f["path"] for f in changes.get("modified", [])}
            
            with st.expander(f"📄 文件列表 · {total_files} 个文件", expanded=False):
                for f in files:
                    icon = _file_icon(f["ext"])
                    file_path = f.get("path", "")
                    
                    # 添加变更标记
                    change_badge = ""
                    if file_path in added_files:
                        change_badge = '<span style="background:#10B981;color:white;font-size:0.65rem;padding:0.1rem 0.4rem;border-radius:999px;margin-left:0.5rem;">新增</span>'
                    elif file_path in modified_files:
                        change_badge = '<span style="background:#F59E0B;color:white;font-size:0.65rem;padding:0.1rem 0.4rem;border-radius:999px;margin-left:0.5rem;">修改</span>'
                    
                    st.markdown(
                        f"""
                        <div class="file-item">
                            <div class="file-item-icon">{icon}</div>
                            <div>
                                <div class="file-item-name">{f['name']}{change_badge}</div>
                                <div class="file-item-meta">{f['relative']} · {_human_size(f['size'])} · {f['modified']}</div>
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
        else:
            st.markdown(
                f"""
                <div class="empty-state">
                    <div class="empty-state-icon">{_svg_icon("inbox", size="2xl", extra_class="icon-glow-soft")}</div>
                    <div class="empty-state-text">未找到最近修改的文件，请尝试增大扫描天数或更换目录</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

    elif not st.session_state.scan_done:
        # Pre-scan empty state
        st.markdown(
            f"""
            <div class="card">
                <div class="card-header">
                    <span class="card-header-icon">{_svg_icon("chart", size="md")}</span>
                    <span class="card-header-title">扫描统计</span>
                </div>
                <div class="empty-state">
                    <div class="empty-state-icon">{_svg_icon("search", size="2xl", extra_class="icon-glow-soft")}</div>
                    <div class="empty-state-text">设置目录并点击扫描，查看文件统计</div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    # -- Quick Config Shortcuts --
    model_name = _get_model_display_name()
    template_name = _get_template_display_name()

    col_ai, col_tpl = st.columns(2)
    with col_ai:
        st.markdown(
            f"""
            <div class="card">
                <div class="card-header">
                    <span class="card-header-icon">{_svg_icon("chip", size="md")}</span>
                    <span class="card-header-title">AI 模型</span>
                </div>
                <div style="padding: 0.6rem 1.2rem 0.2rem;">
                    <span style="font-size:0.92rem;color:#334155;font-weight:500;">{model_name}</span>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        if st.button("⚙️ 修改模型配置", use_container_width=True, key="btn_quick_ai"):
            st.session_state.current_page = "settings"
            st.rerun()

    with col_tpl:
        st.markdown(
            f"""
            <div class="card">
                <div class="card-header">
                    <span class="card-header-icon">{_svg_icon("doc", size="md")}</span>
                    <span class="card-header-title">报告模板</span>
                </div>
                <div style="padding: 0.6rem 1.2rem 0.2rem;">
                    <span style="font-size:0.92rem;color:#334155;font-weight:500;">{template_name}</span>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        if st.button("📋 管理模板", use_container_width=True, key="btn_quick_tpl"):
            st.session_state.current_page = "template"
            st.rerun()

    # -- Generate button & Report --
    if st.session_state.scan_done and files is not None and len(files) > 0:
        # Determine if we have a valid provider
        has_provider = False
        config_mode = st.session_state.get("config_mode", "使用 .env 配置")
        selected_provider_key = st.session_state.get("selected_provider_key")
        manual_provider = None

        if config_mode == "使用 .env 配置":
            has_provider = selected_provider_key is not None
        else:
            base_url = st.session_state.get("manual_base_url", "")
            model_name = st.session_state.get("manual_model_name", "")
            api_key = st.session_state.get("manual_api_key", "")
            if base_url and model_name:
                try:
                    manual_provider = CustomProvider(
                        api_key=api_key,
                        model=model_name,
                        base_url=base_url,
                    )
                    has_provider = True
                except ValueError:
                    has_provider = False

        if has_provider:
            st.markdown('<div class="cta-section">', unsafe_allow_html=True)
            generate_clicked = st.button("🚀 生成工作汇总", type="primary", use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)

            if generate_clicked:
                # Phase 1: Analyze files with progress
                st.markdown(
                    f"""
                    <div class="card">
                        <div class="card-header">
                            <span class="card-header-icon">{_svg_icon("chart", size="md")}</span>
                            <span class="card-header-title">分析进度</span>
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                progress_bar = st.progress(0)
                status_text = st.empty()

                # 使用队列在子线程和主线程之间传递进度
                progress_queue = queue.Queue()

                def _update_progress(current: int, total: int) -> None:
                    """将进度放入队列"""
                    progress_queue.put((current, total))

                try:
                    # Create provider based on config mode
                    if config_mode == "手动配置":
                        provider = manual_provider
                    else:
                        provider = create_provider(selected_provider_key)

                    # 在后台线程中执行分析
                    import threading
                    analysis_result = {"analyses": None, "error": None}

                    # 获取历史快照中的旧文件内容（用于增量分析）
                    _snapshot = get_latest_snapshot()
                    _old_file_contents = get_snapshot_file_contents(_snapshot)
                    logger.info(
                        "从快照中获取了 %d 个文件的旧内容用于增量分析",
                        len(_old_file_contents),
                    )

                    def _run_analysis():
                        try:
                            # 使用增量分析：对比新旧内容，只分析变更部分
                            analyses_list = []
                            total = len(files_to_analyze)
                            for idx, file_info in enumerate(files_to_analyze):
                                file_path = file_info.get("path", "")
                                old_content = _old_file_contents.get(file_path)
                                new_content = read_file(file_path)
                                # 跳过读取失败的文件
                                if new_content.startswith("[Error]") or new_content.startswith("[Skipped]"):
                                    logger.warning("跳过无法读取的文件: %s", file_path)
                                    old_content = None
                                result = analyze_changes(file_info, provider, old_content, new_content)
                                analyses_list.append(result)
                                _update_progress(idx + 1, total)
                            analysis_result["analyses"] = analyses_list
                        except Exception as e:
                            analysis_result["error"] = e

                    analysis_thread = threading.Thread(target=_run_analysis)
                    analysis_thread.start()

                    # 主线程更新进度条
                    current, total_done = 0, 0
                    while analysis_thread.is_alive():
                        try:
                            current, total_done = progress_queue.get(timeout=0.1)
                            pct = current / total_done
                            progress_bar.progress(pct)
                            status_text.text(f"分析中: {current}/{total_done} 个文件")
                        except queue.Empty:
                            pass

                    analysis_thread.join()

                    # 处理结果
                    if analysis_result["error"]:
                        raise analysis_result["error"]

                    analyses = analysis_result["analyses"]
                    progress_bar.progress(1.0)
                    status_text.text(f"分析完成: {len(analyses)} 个文件")

                    # Phase 2: Generate report
                    with st.spinner("正在生成工作汇总..."):
                        week_range = get_week_range(days=scan_days)
                        template_name = st.session_state.selected_template
                        
                        # 获取历史报告用于增量对比
                        previous_report = None
                        latest_snapshot = get_latest_snapshot()
                        if latest_snapshot and latest_snapshot.get("report"):
                            previous_report = latest_snapshot["report"]
                            logger.info("使用历史快照进行增量对比: %s", latest_snapshot.get("id", ""))
                        
                        # 计算内容级文件变更摘要（新增/删除行数）
                        file_changes_line_stats = {}
                        for file_info in files_to_analyze:
                            file_path = file_info.get("path", "")
                            old_content = _old_file_contents.get(file_path)
                            new_content = read_file(file_path)
                            if new_content.startswith("[Error]") or new_content.startswith("[Skipped]"):
                                continue
                            diff = get_file_diff(old_content or "", new_content)
                            if diff["has_changes"]:
                                file_changes_line_stats[file_path] = {
                                    "added_lines": len(diff["added_lines"]),
                                    "removed_lines": len(diff["removed_lines"]),
                                    "change_summary": diff["change_summary"],
                                }
                        
                        # 合并文件级变更信息和行级变更统计
                        file_changes = st.session_state.file_changes or {}
                        file_changes["line_stats"] = file_changes_line_stats
                        
                        report = generate_report(
                            analyses,
                            provider,
                            week_range=week_range,
                            template_name=template_name,
                            previous_report=previous_report,
                            file_changes=file_changes,
                        )

                    st.session_state.report = report
                    st.success("✅ 工作汇总生成完成！")
                    
                    # 显示内容级变更统计
                    if file_changes_line_stats:
                        total_added = sum(s["added_lines"] for s in file_changes_line_stats.values())
                        total_removed = sum(s["removed_lines"] for s in file_changes_line_stats.values())
                        changed_file_count = len(file_changes_line_stats)
                        if total_added > 0 or total_removed > 0:
                            stats_parts = []
                            if total_added > 0:
                                stats_parts.append(f"新增 **{total_added}** 行")
                            if total_removed > 0:
                                stats_parts.append(f"删除 **{total_removed}** 行")
                            stats_text = "，".join(stats_parts)
                            st.info(f"📊 内容变更统计（{changed_file_count} 个文件）：{stats_text}")
                    
                    if previous_report:
                        st.info("📝 已基于历史记录生成增量工作汇总，重点突出新增和修改的内容")
                    
                    # 保存快照到历史记录
                    if st.session_state.save_snapshot_on_generate:
                        snapshot_id = save_snapshot(
                            files=files_to_analyze,
                            analyses=analyses,
                            report=report,
                            metadata={
                                "scan_days": scan_days,
                                "work_dir": work_dir,
                                "template": template_name,
"retention_days": st.session_state.get("retention_days", 30),
                            }
                        )
                        if snapshot_id:
                            logger.info("已保存快照: %s", snapshot_id)

                    # 归档周报到历史记录
                    archive_id = archive_report(
                        report_content=report,
                        date_range=week_range,
                        template_name=template_name,
                        source_files=files_to_analyze,
                        work_dir=work_dir,
                        scan_days=scan_days,
                    )
                    if archive_id:
                        logger.info("周报已归档: %s", archive_id)

                except ValueError as exc:
                    st.error(f"配置错误: {exc}")
                except FileNotFoundError as exc:
                    st.error(f"文件未找到: {exc}")
                except Exception as exc:
                    st.error(f"生成失败: {exc}")
        else:
            st.markdown(
                f"""
                <div class="card">
                    <div style="text-align: center; padding: 1rem 0;">
                        <div style="margin-bottom: 0.5rem;">{_svg_icon("warning", size="xl", extra_class="icon-glow-warning")}</div>
                        <div style="font-size: 0.95rem; color: #475569; font-weight: 500;">
                            请先配置AI模型才能生成工作汇总
                        </div>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            if st.button("⚙️ 前往设置页面配置模型", use_container_width=True):
                st.session_state.current_page = "settings"
                st.rerun()

    # -- Report preview & download --
    if st.session_state.report:
        st.markdown(
            f"""
            <div class="card">
                <div class="card-header">
                    <span class="card-header-icon">{_svg_icon("docDark", size="md")}</span>
                    <span class="card-header-title">生成的工作汇总</span>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.markdown(st.session_state.report)

        st.markdown(
            f"""
            <div class="card">
                <div class="card-header">
                    <span class="card-header-icon">{_svg_icon("download", size="md")}</span>
                    <span class="card-header-title">导出</span>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        col_fmt, col_btn = st.columns([2, 1])
        with col_fmt:
            export_format = st.selectbox(
                "选择导出格式",
                options=["Markdown (.md)", "纯文本 (.txt)", "Word (.docx)", "PDF (.pdf)"],
                index=0,
                help="选择要下载的文件格式",
                label_visibility="collapsed",
            )

        with col_btn:
            # 根据选择的格式生成数据
            week_range = get_week_range(days=scan_days)
            base_name = f"近期工作汇总_{week_range.replace(' ', '').replace('-', '_')}"

            if export_format == "Markdown (.md)":
                download_data = st.session_state.report.encode("utf-8")
                download_filename = f"{base_name}.md"
                download_mime = "text/markdown"
            elif export_format == "纯文本 (.txt)":
                download_data = _md_to_txt(st.session_state.report).encode("utf-8")
                download_filename = f"{base_name}.txt"
                download_mime = "text/plain"
            elif export_format == "Word (.docx)":
                try:
                    download_data = _md_to_docx(st.session_state.report)
                    download_filename = f"{base_name}.docx"
                    download_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                except Exception as e:
                    st.error(f"Word导出失败: {e}")
                    download_data = st.session_state.report.encode("utf-8")
                    download_filename = f"{base_name}.md"
                    download_mime = "text/markdown"
            elif export_format == "PDF (.pdf)":
                try:
                    download_data = _md_to_pdf(st.session_state.report)
                    download_filename = f"{base_name}.pdf"
                    download_mime = "application/pdf"
                except Exception as e:
                    st.error(f"PDF导出失败: {e}")
                    download_data = st.session_state.report.encode("utf-8")
                    download_filename = f"{base_name}.md"
                    download_mime = "text/markdown"
            else:
                download_data = st.session_state.report.encode("utf-8")
                download_filename = f"{base_name}.md"
                download_mime = "text/markdown"

            st.download_button(
                label="📥 下载",
                data=download_data,
                file_name=download_filename,
                mime=download_mime,
                use_container_width=True,
            )


# ---------------------------------------------------------------------------
# Page: Settings (设置页面)
# ---------------------------------------------------------------------------


def _render_test_connection_button(source_key: str, btn_key: str) -> None:
    """Render a test-connection button with status feedback for *source_key*.

    Reads and writes ``st.session_state.source_connection_status[source_key]``.
    """
    status = st.session_state.source_connection_status.get(source_key)
    col_btn, col_status = st.columns([1, 2])

    with col_btn:
        if st.button("🔗 测试连接", key=btn_key, use_container_width=True):
            with st.spinner("正在测试连接..."):
                ok, msg = _test_datasource_connection(source_key)
            st.session_state.source_connection_status[source_key] = {
                "ok": ok,
                "msg": msg,
            }
            st.rerun()

    with col_status:
        if status:
            if status["ok"]:
                st.success(f"✅ {status['msg']}")
            else:
                st.error(f"❌ {status['msg']}")


def template_page():
    """模板管理页面 - 管理内置模板和自定义模板。"""

    # -- Hero Banner --
    _hero_doc = _svg_icon("docStack", size="lg", extra_class="hero-icon icon-glow-white")
    st.markdown(
        f"""
        <div class="hero-banner">
            <div class="hero-badge">{_svg_icon("sparkle", size="xs")} 模板中心</div>
            <h1>{_hero_doc} 模板管理</h1>
            <p>管理内置模板和自定义模板，打造专属周报风格</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # 返回主页按钮
    if st.button("← 返回主页", use_container_width=True):
        st.session_state.current_page = "main"
        st.rerun()

    # -- Tabs --
    tab_builtin, tab_custom = st.tabs(["📦 内置模板", "✏️ 自定义模板"])

    # === Tab 1: Built-in Templates ===
    with tab_builtin:
        st.markdown(
            f"""
            <div class="card">
                <div class="card-header">
                    <span class="card-header-icon">{_svg_icon("docStack", size="md")}</span>
                    <span class="card-header-title">内置模板</span>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        all_templates = list_templates()
        builtin_keys = ["standard", "concise", "detailed"]
        builtin_templates = {
            k: v for k, v in all_templates.items() if k in builtin_keys
        }

        # Template metadata for display
        template_meta = {
            "standard": {"icon": "📄", "label": "标准模板", "color": "#3b82f6"},
            "concise": {"icon": "📝", "label": "简洁模板", "color": "#10b981"},
            "detailed": {"icon": "📚", "label": "详细模板", "color": "#8b5cf6"},
        }

        if builtin_templates:
            for key, data in builtin_templates.items():
                meta = template_meta.get(key, {"icon": "📄", "label": key, "color": "#6b7280"})
                with st.expander(f"{meta['icon']} {data['name']}", expanded=False):
                    # Description
                    st.markdown(f"**描述：** {data.get('description', '暂无描述')}")

                    # Sections
                    sections = data.get("sections", [])
                    if sections:
                        sections_text = "、".join(sections)
                        st.markdown(f"**包含章节：** {sections_text}")

                    # Created date
                    created_at = data.get("created_at", "未知")
                    st.caption(f"创建日期：{created_at}")

                    # Preview button
                    if st.button(f"👁️ 预览模板内容", key=f"preview_{key}", use_container_width=True):
                        st.session_state[f"show_preview_{key}"] = True

                    # Show preview if toggled
                    if st.session_state.get(f"show_preview_{key}", False):
                        st.divider()
                        st.markdown("**模板内容预览：**")
                        template_content = data.get("template", "")
                        st.code(template_content, language="markdown")
                        if st.button("收起预览", key=f"hide_preview_{key}"):
                            st.session_state[f"show_preview_{key}"] = False
                            st.rerun()

                    # Select as active template
                    current_tpl = st.session_state.get("selected_template")
                    is_selected = current_tpl == key
                    if is_selected:
                        st.success("✅ 当前正在使用此模板")
                    else:
                        if st.button(f"📌 使用此模板", key=f"use_{key}", use_container_width=True):
                            st.session_state.selected_template = key
                            st.success(f"已切换到 {data['name']}")
                            st.rerun()
        else:
            st.warning("未找到内置模板，请检查 templates/builtin 目录")

    # === Tab 2: Custom Templates ===
    with tab_custom:
        st.markdown(
            f"""
            <div class="card">
                <div class="card-header">
                    <span class="card-header-icon">{_svg_icon("docDark", size="md")}</span>
                    <span class="card-header-title">自定义模板</span>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        # -- Section 1: Upload and Extract Template --
        st.markdown(
            f'<div class="settings-section-title">{_svg_icon("folder", size="md")} 从文件提取模板</div>',
            unsafe_allow_html=True,
        )

        uploaded_file = st.file_uploader(
            "上传周报文件",
            type=["md", "txt", "docx", "pdf"],
            help="上传一份现有的周报，系统将自动提取模板结构（支持 .md, .txt, .docx, .pdf 格式）",
            key="template_upload",
        )

        if uploaded_file is not None:
            # Read uploaded file content based on file type
            file_ext = Path(uploaded_file.name).suffix.lower()

            if file_ext in ['.md', '.txt']:
                file_content = uploaded_file.read().decode("utf-8")
            elif file_ext in ['.docx', '.pdf']:
                import tempfile
                suffix = file_ext
                with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                    tmp.write(uploaded_file.read())
                    tmp_path = tmp.name
                from core.template_manager import read_file_content
                file_content = read_file_content(tmp_path)
                Path(tmp_path).unlink()
            else:
                st.error("不支持的文件格式")
                file_content = ""

            if file_content:
                # Show preview
                with st.expander("📄 预览上传内容", expanded=False):
                    preview_text = file_content[:500] + "..." if len(file_content) > 500 else file_content
                    st.text(preview_text)

                template_name_input = st.text_input(
                    "模板名称",
                    value="",
                    placeholder="输入模板名称，如：我的周报模板",
                    help="为提取的模板命名",
                    key="extract_template_name",
                )

                # LLM analysis option
                use_llm_analysis = st.toggle(
                    "🤖 使用 AI 智能分析模板",
                    value=True,
                    help="使用 AI 分析模板结构，生成更优化的提示词模板",
                    key="use_llm_toggle",
                )

                if st.button("📥 提取模板", use_container_width=True, key="extract_btn"):
                    if not template_name_input.strip():
                        st.error("请输入模板名称")
                    else:
                        with st.spinner("正在提取模板..."):
                            if use_llm_analysis:
                                # 使用 LLM 分析模板
                                from core.template_manager import analyze_template_with_llm

                                # 获取 LLM provider
                                provider = None
                                config_mode = st.session_state.get("config_mode", "使用 .env 配置")
                                selected_provider_key = st.session_state.get("selected_provider_key")

                                if config_mode == "使用 .env 配置" and selected_provider_key:
                                    try:
                                        provider = create_provider(selected_provider_key)
                                    except Exception as e:
                                        st.warning(f"无法创建 LLM provider: {e}，将使用基础提取")
                                elif config_mode == "手动配置":
                                    base_url = st.session_state.get("manual_base_url", "")
                                    api_key = st.session_state.get("manual_api_key", "")
                                    model_name = st.session_state.get("manual_model_name", "")
                                    if base_url and model_name:
                                        try:
                                            provider = CustomProvider(
                                                api_key=api_key,
                                                model=model_name,
                                                base_url=base_url,
                                            )
                                        except Exception as e:
                                            st.warning(f"无法创建 LLM provider: {e}，将使用基础提取")

                                if provider:
                                    template_data = analyze_template_with_llm(
                                        file_content, provider, template_name_input.strip()
                                    )
                                    st.info("✨ AI 已分析模板结构并生成优化的提示词")
                                else:
                                    # 回退到基础提取
                                    template_data = extract_template_from_report(
                                        file_content, template_name_input.strip()
                                    )
                            else:
                                # 使用基础提取
                                template_data = extract_template_from_report(
                                    file_content, template_name_input.strip()
                                )

                            success = save_custom_template(
                                template_name_input.strip(), template_data
                            )

                        if success:
                            st.success(f"模板 '{template_name_input}' 提取并保存成功！")

                            # 显示分析结果
                            if use_llm_analysis and "analysis" in template_data:
                                with st.expander("📊 AI 分析结果", expanded=True):
                                    st.markdown(template_data["analysis"])

                            st.rerun()
                        else:
                            st.error("模板保存失败")

        st.divider()

        # -- Section 2: Existing Custom Templates --
        st.markdown(
            f'<div class="settings-section-title">{_svg_icon("docDark", size="md")} 已保存的自定义模板</div>',
            unsafe_allow_html=True,
        )

        all_templates = list_templates()
        custom_templates = {
            k: v for k, v in all_templates.items()
            if k not in ["standard", "concise", "detailed"]
        }

        if custom_templates:
            for key, data in custom_templates.items():
                with st.expander(f"📝 {data['name']}", expanded=False):
                    # Description
                    description = data.get("description", "暂无描述")
                    st.markdown(f"**描述：** {description}")

                    # Sections
                    sections = data.get("sections", [])
                    if sections:
                        sections_text = "、".join(sections)
                        st.markdown(f"**包含章节：** {sections_text}")

                    # Created date
                    created_at = data.get("created_at", "未知")
                    st.caption(f"创建日期：{created_at}")

                    # Preview
                    template_content = data.get("template", "")
                    if template_content:
                        st.markdown("**👁️ 预览模板内容：**")
                        st.code(template_content, language="markdown")

                    # Action buttons
                    col_use, col_delete = st.columns(2)
                    with col_use:
                        current_tpl = st.session_state.get("selected_template")
                        is_selected = current_tpl == key
                        if is_selected:
                            st.success("✅ 当前使用中")
                        else:
                            if st.button(f"📌 使用此模板", key=f"use_custom_{key}", use_container_width=True):
                                st.session_state.selected_template = key
                                st.success(f"已切换到 {data['name']}")
                                st.rerun()
                    with col_delete:
                        if st.button(f"🗑️ 删除", key=f"del_custom_{key}", use_container_width=True):
                            delete_custom_template(key)
                            st.success(f"已删除模板 '{data['name']}'")
                            st.rerun()
        else:
            st.info("暂无自定义模板，请从文件提取或上传创建")


def history_page():
    """历史记录页面 - 周报归档查看与管理。"""
    
    # -- Hero Banner --
    _hero_clock = _svg_icon("clock", size="lg", extra_class="hero-icon icon-glow-white")
    st.markdown(
        f"""
        <div class="hero-banner">
            <div class="hero-badge">{_svg_icon("sparkle", size="xs")} 历史归档</div>
            <h1>{_hero_clock} 历史记录</h1>
            <p>查看所有已生成的周报归档，点击查看详情</p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    
    # 返回主页按钮
    if st.button("← 返回主页", use_container_width=True):
        st.session_state.current_page = "main"
        st.rerun()
    
    # ── 初始化 session state ──────────────────────────────────────
    if "history_view_report_id" not in st.session_state:
        st.session_state.history_view_report_id = None
    
    # ── 概览统计 ──────────────────────────────────────────────────
    stats = get_archive_stats()
    _total = stats.get("total_count", 0)
    _latest = stats.get("latest_time")
    _oldest = stats.get("oldest_time")
    
    try:
        _latest_display = _latest[:16].replace("T", " ") if _latest else "无"
    except Exception:
        _latest_display = str(_latest) if _latest else "无"
    try:
        _oldest_display = _oldest[:16].replace("T", " ") if _oldest else "无"
    except Exception:
        _oldest_display = str(_oldest) if _oldest else "无"
    
    st.markdown(
        f"""
        <div class="card">
            <div class="card-header">
                <span class="card-header-icon">{_svg_icon("chart", size="md")}</span>
                <span class="card-header-title">归档概览</span>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    
    _m1, _m2, _m3 = st.columns(3)
    with _m1:
        st.metric("总周报数", _total)
    with _m2:
        st.metric("最近生成", _latest_display)
    with _m3:
        st.metric("最早记录", _oldest_display)
    
    # ── 如果正在查看某条周报详情 ──────────────────────────────────
    if st.session_state.history_view_report_id:
        _show_report_detail(st.session_state.history_view_report_id)
        return
    
    # ── 周报归档列表 ──────────────────────────────────────────────
    st.markdown(
        f"""
        <div class="card" style="margin-top:1rem;">
            <div class="card-header">
                <span class="card-header-icon">{_svg_icon("doc", size="md")}</span>
                <span class="card-header-title">周报归档列表</span>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    
    # 获取归档列表
    reports = list_archived_reports(limit=50)
    
    if not reports:
        st.markdown(
            f"""
            <div class="empty-state">
                <div class="empty-state-icon">{_svg_icon("inbox", size="2xl", extra_class="icon-glow-soft")}</div>
                <div class="empty-state-text">暂无周报归档</div>
                <div style="font-size:0.85rem;color:#64748B;margin-top:0.5rem;">
                    生成周报后会自动归档保存，可在此处查看历史周报
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        return
    
    # 显示周报列表
    for idx, report_item in enumerate(reports):
        report_id = report_item.get("report_id", "")
        created_at = report_item.get("created_at", "")
        date_range = report_item.get("date_range", "")
        template_name = report_item.get("template_name", "")
        summary = report_item.get("summary", "")
        source_files_count = report_item.get("source_files_count", 0)
        
        # 格式化时间显示
        try:
            dt = datetime.fromisoformat(created_at)
            time_display = dt.strftime("%Y-%m-%d %H:%M:%S")
            relative_time = _get_relative_time(dt)
        except Exception:
            time_display = created_at[:19] if created_at else "未知"
            relative_time = ""
        
        # 构建标题
        title_parts = []
        if date_range:
            title_parts.append(f"📅 {date_range}")
        title_parts.append(f"🕐 {time_display}")
        if relative_time:
            title_parts.append(f"({relative_time})")
        
        title = " · ".join(title_parts)
        
        # 使用 expander 展示列表项
        with st.expander(title, expanded=False):
            # 摘要信息
            st.markdown(
                f"""
                <div style="display:grid;grid-template-columns:repeat(auto-fit,minmax(140px,1fr));gap:0.75rem;margin-bottom:1rem;">
                    <div style="background:rgba(255,255,255,0.65);padding:0.6rem;border-radius:8px;">
                        <div style="font-size:0.72rem;color:#64748B;text-transform:uppercase;">日期范围</div>
                        <div style="font-size:0.95rem;font-weight:600;color:#1E293B;">{date_range or '未知'}</div>
                    </div>
                    <div style="background:rgba(255,255,255,0.65);padding:0.6rem;border-radius:8px;">
                        <div style="font-size:0.72rem;color:#64748B;text-transform:uppercase;">使用模板</div>
                        <div style="font-size:0.95rem;font-weight:600;color:#1E293B;">{template_name or '默认'}</div>
                    </div>
                    <div style="background:rgba(255,255,255,0.65);padding:0.6rem;border-radius:8px;">
                        <div style="font-size:0.72rem;color:#64748B;text-transform:uppercase;">来源文件</div>
                        <div style="font-size:0.95rem;font-weight:600;color:#1E293B;">{source_files_count} 个</div>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            
            # 摘要文本
            if summary:
                st.markdown(f"**摘要：** {summary}")
            
            # 操作按钮
            col_view, col_delete = st.columns([2, 1])
            with col_view:
                if st.button(
                    "📖 查看详情",
                    key=f"view_report_{report_id}",
                    use_container_width=True,
                    type="primary",
                ):
                    st.session_state.history_view_report_id = report_id
                    st.rerun()
            with col_delete:
                if st.button(
                    "🗑️ 删除",
                    key=f"del_report_{report_id}",
                    use_container_width=True,
                ):
                    if delete_archived_report(report_id):
                        st.success("已删除")
                        st.rerun()
                    else:
                        st.error("删除失败")


def _show_report_detail(report_id: str):
    """显示单条周报归档的详细内容。"""
    
    # 返回列表按钮
    if st.button("← 返回列表", key="back_to_list", use_container_width=True):
        st.session_state.history_view_report_id = None
        st.rerun()
    
    # 获取周报详情
    report_data = get_archived_report(report_id)
    
    if not report_data:
        st.error("未找到该周报记录，可能已被删除。")
        st.session_state.history_view_report_id = None
        return
    
    # 解析数据
    created_at = report_data.get("created_at", "")
    date_range = report_data.get("date_range", "")
    template_name = report_data.get("template_name", "")
    report_content = report_data.get("report_content", "")
    source_files_count = report_data.get("source_files_count", 0)
    source_files_info = report_data.get("source_files_info", [])
    work_dir = report_data.get("work_dir", "")
    scan_days = report_data.get("scan_days", 7)
    export_path = report_data.get("export_path", "")
    metadata = report_data.get("metadata", {})
    
    # 格式化时间
    try:
        dt = datetime.fromisoformat(created_at)
        time_display = dt.strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        time_display = created_at[:19] if created_at else "未知"
    
    # ── 元信息卡片 ──────────────────────────────────────────────
    _icon_doc = _svg_icon("doc", size="md")
    st.markdown(
        f"""
        <div class="card">
            <div class="card-header">
                <span class="card-header-icon">{_icon_doc}</span>
                <span class="card-header-title">周报详情</span>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    
    # 元信息网格
    auto_tag = " 🤖 自动生成" if metadata.get("auto_generated") else ""
    st.markdown(
        f"""
        <div style="display:grid;grid-template-columns:repeat(auto-fit,minmax(160px,1fr));gap:0.75rem;margin-bottom:1.25rem;">
            <div style="background:rgba(255,255,255,0.65);padding:0.75rem;border-radius:8px;">
                <div style="font-size:0.72rem;color:#64748B;text-transform:uppercase;">生成时间</div>
                <div style="font-size:0.9rem;font-weight:600;color:#1E293B;">{time_display}{auto_tag}</div>
            </div>
            <div style="background:rgba(255,255,255,0.65);padding:0.75rem;border-radius:8px;">
                <div style="font-size:0.72rem;color:#64748B;text-transform:uppercase;">日期范围</div>
                <div style="font-size:0.9rem;font-weight:600;color:#1E293B;">{date_range or '未指定'}</div>
            </div>
            <div style="background:rgba(255,255,255,0.65);padding:0.75rem;border-radius:8px;">
                <div style="font-size:0.72rem;color:#64748B;text-transform:uppercase;">使用模板</div>
                <div style="font-size:0.9rem;font-weight:600;color:#1E293B;">{template_name or '默认'}</div>
            </div>
            <div style="background:rgba(255,255,255,0.65);padding:0.75rem;border-radius:8px;">
                <div style="font-size:0.72rem;color:#64748B;text-transform:uppercase;">来源文件</div>
                <div style="font-size:0.9rem;font-weight:600;color:#1E293B;">{source_files_count} 个</div>
            </div>
            <div style="background:rgba(255,255,255,0.65);padding:0.75rem;border-radius:8px;">
                <div style="font-size:0.72rem;color:#64748B;text-transform:uppercase;">工作目录</div>
                <div style="font-size:0.9rem;font-weight:600;color:#1E293B;">{work_dir or '未记录'}</div>
            </div>
            <div style="background:rgba(255,255,255,0.65);padding:0.75rem;border-radius:8px;">
                <div style="font-size:0.72rem;color:#64748B;text-transform:uppercase;">扫描天数</div>
                <div style="font-size:0.9rem;font-weight:600;color:#1E293B;">{scan_days} 天</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    
    # ── 周报正文 ──────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 📄 周报正文")
    
    if report_content:
        st.markdown(report_content)
    else:
        st.warning("该记录没有保存周报内容。")
    
    # ── 来源文件列表 ──────────────────────────────────────────────
    if source_files_info and isinstance(source_files_info, list) and len(source_files_info) > 0:
        st.markdown("---")
        with st.expander(f"📁 来源文件列表（{source_files_count} 个）", expanded=False):
            for f_info in source_files_info:
                fname = f_info.get("name", "")
                frel = f_info.get("relative", "")
                fext = f_info.get("ext", "")
                fsize = f_info.get("size", 0)
                fmod = f_info.get("modified", "")
                
                # 格式化文件大小
                if fsize > 1024 * 1024:
                    size_str = f"{fsize / 1024 / 1024:.1f} MB"
                elif fsize > 1024:
                    size_str = f"{fsize / 1024:.1f} KB"
                else:
                    size_str = f"{fsize} B"
                
                st.markdown(f"- **{fname}** ({frel}) — {size_str} — 修改于 {fmod}")
    
    # ── 导出操作 ──────────────────────────────────────────────────
    st.markdown("---")
    col_export_md, col_export_txt = st.columns(2)
    
    with col_export_md:
        if report_content:
            st.download_button(
                label="📥 导出 Markdown",
                data=report_content,
                file_name=f"周报_{date_range.replace(' ', '').replace('-', '_') if date_range else time_display[:10]}.md",
                mime="text/markdown",
                use_container_width=True,
            )
    
    with col_export_txt:
        if report_content:
            st.download_button(
                label="📥 导出纯文本",
                data=report_content,
                file_name=f"周报_{date_range.replace(' ', '').replace('-', '_') if date_range else time_display[:10]}.txt",
                mime="text/plain",
                use_container_width=True,
            )


def _get_relative_time(dt: datetime) -> str:
    """获取相对时间描述。"""
    now = datetime.now()
    diff = now - dt
    
    if diff.days > 30:
        return f"{diff.days // 30} 个月前"
    elif diff.days > 0:
        return f"{diff.days} 天前"
    elif diff.seconds > 3600:
        return f"{diff.seconds // 3600} 小时前"
    elif diff.seconds > 60:
        return f"{diff.seconds // 60} 分钟前"
    else:
        return "刚刚"


def settings_page():
    """设置页面 - AI模型/分析参数/系统."""

    # -- Hero Banner for Settings --
    _hero_gear = _svg_icon("gear", size="lg", extra_class="hero-icon icon-glow-white icon-rotate")
    st.markdown(
        f"""
        <div class="hero-banner">
            <div class="hero-badge">{_svg_icon("sparkle", size="xs")} 配置中心</div>
            <h1>{_hero_gear} 设置</h1>
            <p>配置AI模型、分析参数和系统选项</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # -- Tabs for settings sections --
    tab_model, tab_analysis, tab_auto_scan, tab_system = st.tabs([
        "🤖 AI模型",
        "📊 分析参数",
        "🔄 自动扫描",
        "⚙️ 系统",
    ])

    # === Tab 1: AI模型 ===
    with tab_model:
        st.markdown(
            f'<div class="settings-section-title">{_svg_icon("chip", size="md")} AI模型配置</div>',
            unsafe_allow_html=True,
        )

        config_mode = st.selectbox(
            "配置模式",
            options=["使用 .env 配置", "手动配置"],
            index=0 if st.session_state.config_mode == "使用 .env 配置" else 1,
            help="选择模型配置来源：使用环境变量文件或手动输入",
        )
        st.session_state.config_mode = config_mode

        if config_mode == "使用 .env 配置":
            # -- Model selection from .env --
            provider_display_list, display_to_key = _get_provider_choices()

            if provider_display_list:
                # Determine current index
                current_key = st.session_state.get("selected_provider_key")
                current_idx = 0
                if current_key:
                    for i, display in enumerate(provider_display_list):
                        if display_to_key[display] == current_key:
                            current_idx = i
                            break

                selected_display = st.selectbox(
                    "🤖 模型选择",
                    options=provider_display_list,
                    index=current_idx,
                    help="自动检测已配置API Key的模型",
                )
                st.session_state.selected_provider_key = display_to_key[selected_display]
                st.success(f"已检测到 {len(provider_display_list)} 个可用模型")
            else:
                st.session_state.selected_provider_key = None
                st.warning("未检测到已配置的API Key")
                st.caption("请在 `.env` 文件中配置至少一个模型的API Key")

        else:
            # -- Manual configuration --
            manual_base_url = st.text_input(
                "🌐 Base URL",
                value=st.session_state.manual_base_url,
                placeholder="https://api.example.com/v1",
                help="OpenAI兼容API的基础URL",
            )
            st.session_state.manual_base_url = manual_base_url

            manual_api_key = st.text_input(
                "🔑 API Key",
                value=st.session_state.manual_api_key,
                type="password",
                placeholder="sk-...",
                help="API密钥（可选，某些本地服务不需要）",
            )
            st.session_state.manual_api_key = manual_api_key

            manual_model_name = st.text_input(
                "📦 模型名称",
                value=st.session_state.manual_model_name,
                placeholder="gpt-4o-mini",
                help="要使用的模型名称",
            )
            st.session_state.manual_model_name = manual_model_name

            if manual_base_url and manual_model_name:
                try:
                    CustomProvider(
                        api_key=manual_api_key,
                        model=manual_model_name,
                        base_url=manual_base_url,
                    )
                    st.success("手动配置已就绪")
                except ValueError as e:
                    st.error(f"配置错误: {e}")
            else:
                st.caption("请输入 Base URL 和模型名称")

        # -- Test connection button --
        st.divider()
        if st.button("🔗 测试连接", use_container_width=True):
            provider = None
            try:
                if config_mode == "手动配置":
                    base_url = st.session_state.get("manual_base_url", "")
                    api_key = st.session_state.get("manual_api_key", "")
                    model_name = st.session_state.get("manual_model_name", "")
                    if base_url and model_name:
                        provider = CustomProvider(
                            api_key=api_key,
                            model=model_name,
                            base_url=base_url,
                        )
                    else:
                        st.error("请先填写完整配置")
                        return
                else:
                    provider_key = st.session_state.get("selected_provider_key")
                    if provider_key:
                        provider = create_provider(provider_key)
                    else:
                        st.error("请先选择模型")
                        return

                with st.spinner("正在测试连接..."):
                    # Simple test: generate a short text
                    messages = [{"role": "user", "content": "请回复'连接成功'两个字"}]
                    result = provider.chat_completion(messages, max_tokens=50)
                    st.success(f"连接成功！模型响应: {result[:100]}")
            except Exception as e:
                st.error(f"连接失败: {e}")

    # === Tab 2: 分析参数 ===
    with tab_analysis:
        st.markdown(
            f'<div class="settings-section-title">{_svg_icon("chart", size="md")} 分析参数</div>',
            unsafe_allow_html=True,
        )

        concurrency = st.number_input(
            "⚡ 并发数",
            min_value=1,
            max_value=10,
            value=st.session_state.concurrency,
            help="同时分析的文件数量（较高的并发数可能消耗更多API额度）",
            key="concurrency_widget",
            on_change=lambda: setattr(st.session_state, 'concurrency', st.session_state.concurrency_widget),
        )

        use_cache = st.toggle(
            "💾 启用缓存",
            value=st.session_state.use_cache,
            help="缓存已分析的文件结果，避免重复调用API",
        )
        st.session_state.use_cache = use_cache

        log_level = st.selectbox(
            "📋 日志级别",
            options=["DEBUG", "INFO", "WARNING", "ERROR"],
            index=["DEBUG", "INFO", "WARNING", "ERROR"].index(st.session_state.log_level),
            help="设置日志输出级别",
        )
        st.session_state.log_level = log_level

    # === Tab 3: 自动扫描 ===
    with tab_auto_scan:
        st.markdown(
            f'<div class="settings-section-title">{_svg_icon("clock", size="md")} 自动扫描设置</div>',
            unsafe_allow_html=True,
        )

        auto_scan_enabled = st.toggle(
            "🔄 启用每日自动扫描",
            value=st.session_state.get("auto_scan_enabled", False),
            help="开启后将每天自动扫描文件并生成工作汇总",
        )
        st.session_state.auto_scan_enabled = auto_scan_enabled

        if auto_scan_enabled:
            st.markdown("**⏰ 执行时间设置**")

            col_hour, col_minute = st.columns(2)
            with col_hour:
                auto_scan_hour = st.number_input(
                    "小时",
                    min_value=0,
                    max_value=23,
                    value=st.session_state.get("auto_scan_hour", 9),
                    help="设置执行时间的小时（0-23）",
                )
                st.session_state.auto_scan_hour = auto_scan_hour

            with col_minute:
                auto_scan_minute = st.number_input(
                    "分钟",
                    min_value=0,
                    max_value=59,
                    value=st.session_state.get("auto_scan_minute", 0),
                    help="设置执行时间的分钟（0-59）",
                )
                st.session_state.auto_scan_minute = auto_scan_minute

            time_str = f"{auto_scan_hour:02d}:{auto_scan_minute:02d}"
            st.info(f"✅ 自动扫描已启用，将在每天 {time_str} 执行")

            # 更新调度器配置并启动
            scheduler = st.session_state.auto_scheduler
            scheduler.set_callback(auto_scan_callback)
            scheduler.set_schedule(
                enabled=True,
                time_str=time_str,
            )
            # 保存扫描配置到文件（供后台线程使用，因为线程无法访问 session_state）
            scheduler.save_scan_config({
                "work_dir": st.session_state.get("work_dir", ""),
                "scan_days": st.session_state.get("scan_days", 7),
                "selected_template": st.session_state.get("selected_template"),
                "config_mode": st.session_state.get("config_mode", "使用 .env 配置"),
                "selected_provider_key": st.session_state.get("selected_provider_key"),
                "manual_base_url": st.session_state.get("manual_base_url", ""),
                "manual_api_key": st.session_state.get("manual_api_key", ""),
                "manual_model_name": st.session_state.get("manual_model_name", ""),
                "retention_days": st.session_state.get("retention_days", 30),
            })
        else:
            st.info("ℹ️ 自动扫描未启用")

            # 停止调度器
            scheduler = st.session_state.auto_scheduler
            scheduler.set_schedule(enabled=False, time_str="09:00")

    # === Tab 4: 系统 ===
    with tab_system:
        st.markdown(
            f'<div class="settings-section-title">{_svg_icon("gear", size="md")} 系统设置</div>',
            unsafe_allow_html=True,
        )

        # 快照设置
        st.markdown(
            f'<div class="settings-section-title">{_svg_icon("clock", size="md")} 快照与历史</div>',
            unsafe_allow_html=True,
        )

        save_snapshot_on_generate = st.toggle(
            "📸 生成报告时自动保存快照",
            value=st.session_state.save_snapshot_on_generate,
            help="每次生成报告时自动保存文件快照，用于变更追踪",
        )
        st.session_state.save_snapshot_on_generate = save_snapshot_on_generate

        retention_days = st.number_input(
            "💾 快照保留天数",
            min_value=1,
            max_value=365,
            value=st.session_state.get("retention_days", 30),
            help="超过保留天数的历史快照将自动删除",
            key="retention_days_widget",
            on_change=lambda: setattr(st.session_state, 'retention_days', st.session_state.retention_days_widget),
        )

        # 显示历史记录摘要
        timeline_summary = get_timeline_summary()

        if timeline_summary["total_snapshots"] > 0:
            st.markdown(
                f"""
                <div style="background:rgba(255,255,255,0.65);border-radius:12px;padding:1rem;margin-top:0.5rem;">
                    <div style="font-size:0.85rem;color:#475569;">
                        📊 历史记录：共 {timeline_summary["total_snapshots"]} 个快照
                    </div>
                    <div style="font-size:0.78rem;color:#64748B;margin-top:0.25rem;">
                        最近：{timeline_summary["latest_time"][:19] if timeline_summary["latest_time"] else "无"}
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )

            if st.button("📜 查看历史记录", use_container_width=True):
                st.session_state.current_page = "history"
                st.rerun()
        else:
            st.info("📝 暂无历史记录，生成报告时会自动保存")

        st.divider()

        # 导出配置
        st.markdown(
            f'<div class="settings-section-title">{_svg_icon("externalLink", size="md")} 导出配置</div>',
            unsafe_allow_html=True,
        )

        _export_snapshots = list_snapshots(limit=100)
        if _export_snapshots:
            _export_json = json.dumps(_export_snapshots, ensure_ascii=False, indent=2, default=str)
            st.download_button(
                label="📥 导出全部快照",
                data=_export_json,
                file_name=f"snapshots_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
                use_container_width=True,
            )
        else:
            st.button("📥 导出全部快照", disabled=True, use_container_width=True)
            st.caption("暂无可导出的快照")

    # -- Bottom actions --
    st.divider()
    col_save, col_back = st.columns(2)
    with col_save:
        if st.button("💾 保存配置", use_container_width=True, type="primary"):
            st.success("配置已保存")

    with col_back:
        if st.button("🏠 返回主页", use_container_width=True):
            st.session_state.current_page = "main"
            st.rerun()


# ---------------------------------------------------------------------------
# Navigation: Main page (default) + Settings page
# ---------------------------------------------------------------------------

# Use tabs-based navigation since st.navigation requires separate page files
# We use session_state to track current page
if "current_page" not in st.session_state:
    st.session_state.current_page = "main"

# Sidebar navigation
with st.sidebar:
    _sidebar_compass = _svg_icon("compass", size="xl", extra_class="icon-shimmer")
    st.markdown(
        f"""
        <div class="sidebar-brand">
            <span class="sidebar-brand-icon">{_sidebar_compass}</span>
            <div class="sidebar-brand-text">周报终结者</div>
            <div class="sidebar-brand-sub">Weekly Report Agent</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown('<div class="sidebar-section-label">导航</div>', unsafe_allow_html=True)

    main_btn = st.button("📝 主页", use_container_width=True, type="primary" if st.session_state.current_page == "main" else "secondary")
    template_btn = st.button("📋 模板", use_container_width=True, type="primary" if st.session_state.current_page == "template" else "secondary")
    history_btn = st.button("🕐 历史", use_container_width=True, type="primary" if st.session_state.current_page == "history" else "secondary")
    settings_btn = st.button("⚙️ 设置", use_container_width=True, type="primary" if st.session_state.current_page == "settings" else "secondary")

    # 处理按钮点击
    if main_btn and st.session_state.current_page != "main":
        st.session_state.current_page = "main"
        st.rerun()
    elif template_btn and st.session_state.current_page != "template":
        st.session_state.current_page = "template"
        st.rerun()
    elif history_btn and st.session_state.current_page != "history":
        st.session_state.current_page = "history"
        st.rerun()
    elif settings_btn and st.session_state.current_page != "settings":
        st.session_state.current_page = "settings"
        st.rerun()

    st.divider()

    # -- Data source status in sidebar --
    st.markdown('<div class="sidebar-section-label">数据源状态</div>', unsafe_allow_html=True)

    # Local files status
    _sb_files = st.session_state.files
    _sb_file_count = len(_sb_files) if (st.session_state.scan_done and _sb_files is not None) else 0
    _sb_file_icon = "✅" if _sb_file_count > 0 else "⚪"
    st.markdown(
        f'<div style="font-size:0.78rem;color:rgba(255,255,255,0.7);padding:0.25rem 0;">'
        f'{_sb_file_icon} 本地文件: {_sb_file_count} 个</div>',
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="sidebar-section-label" style="margin-top: 0.5rem;">快速指南</div>
        <div class="tip-box" style="margin-top: 0.5rem;">
            <div class="tip-box-content">
                ① 输入工作目录路径<br>
                ② 选择扫描天数<br>
                ③ 配置AI模型<br>
                ④ 点击 <strong>生成工作汇总</strong><br>
                ⑤ 等待分析完成<br>
                ⑥ 选择格式并下载
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

# Render current page
if st.session_state.current_page == "template":
    template_page()
elif st.session_state.current_page == "history":
    history_page()
elif st.session_state.current_page == "settings":
    settings_page()
else:
    main_page()
